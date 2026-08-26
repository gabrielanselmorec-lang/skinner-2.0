"""Testes das regras de montagem do DOCX PEI."""
import os
import sys
import warnings
from datetime import date

import pandas as pd

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from docx import Document

from app.services.pei_docx import (
    aplicar_fonte_pei,
    gerar_doc_completo,
    gerar_resumo_trimestral_local,
    gerar_resumo_comportamentos_problema_local,
    preparar_programas_pei,
    resumo_alvos_por_objetivo,
)


def test_resumo_alvos_por_objetivo_exibe_codigo_e_nome_do_objetivo():
    areas = {1: [{"programa": "Mando com frases", "texto": "Usar frases para pedir."}]}
    df_hist = pd.DataFrame({
        "programa": ["Mando com frases"],
        "date": ["2026-07-01"],
        "independent_rate": [80],
        "phase": [""],
    })
    df_alvos = pd.DataFrame({
        "programa": ["Mando com frases"],
        "target_name": ["Pedir ajuda"],
        "date": ["2026-07-01"],
        "independent_rate": [80],
        "phase": [""],
    })

    resultado = resumo_alvos_por_objetivo(
        areas,
        df_hist,
        df_alvos,
        df_alvos,
        pd.DataFrame(),
        date(2026, 7, 1),
        date(2026, 7, 31),
    )

    assert resultado.iloc[0]["Objetivo"] == "1.1 Mando com frases"


def test_preparar_programas_pei_usa_template_da_biblioteca_quando_objetivo_vazio():
    df_prog = pd.DataFrame({"programa": ["Mando com frases"], "objective": [""]})
    df_lib = pd.DataFrame({
        "name": ["Mando com frases"],
        "mastery_threshold_percent": [90],
        "objective_template": ["Emitir mandos com frases funcionais."],
    })

    resultado = preparar_programas_pei(df_prog, df_lib)

    assert resultado.iloc[0]["objective"] == "Emitir mandos com frases funcionais."


def test_aplicar_fonte_pei_nao_usa_lookup_de_style_id():
    doc = Document()

    with warnings.catch_warnings(record=True) as capturados:
        warnings.simplefilter("always")
        aplicar_fonte_pei(doc)

    mensagens = [str(item.message) for item in capturados]
    assert not any("style lookup by style_id is deprecated" in msg for msg in mensagens)


def test_resumo_clinico_inclui_objetivos_alvos_e_evolucoes_no_docx():
    df_hist = pd.DataFrame({
        "programa": ["Mando com frases", "Mando com frases"],
        "date": ["2026-07-01", "2026-07-15"],
        "date_pd": pd.to_datetime(["2026-07-01", "2026-07-15"]),
        "independent_rate": [50, 80],
        "prompt_rate": [50, 20],
        "phase": ["", ""],
        "objective": ["Usar frases para pedir.", "Usar frases para pedir."],
        "evolution": ["Paciente solicitou itens com frase curta.", "Paciente aumentou espontaneidade nos pedidos."],
    })
    df_alvos = pd.DataFrame({
        "programa": ["Mando com frases"],
        "target_name": ["Pedir ajuda"],
        "date": ["2026-07-15"],
        "date_pd": pd.to_datetime(["2026-07-15"]),
        "independent_rate": [80],
        "phase": [""],
        "evolution": ["Usou frase para solicitar ajuda durante a atividade."],
    })
    df_beh = pd.DataFrame({
        "comportamento": ["Fuga de demanda"],
        "date": ["2026-07-15"],
        "count": [1],
        "rate": [0.2],
        "evolution": ["Retornou para a atividade apos redirecionamento."],
    })

    resumo = gerar_resumo_trimestral_local(
        df_hist,
        df_alvos,
        df_beh,
        "Media geral dos objetivos",
        date(2026, 7, 1),
        date(2026, 7, 31),
    )
    assert "Desempenho por objetivo" in resumo
    assert "Alvos com melhor desempenho" in resumo
    assert "tendência" in resumo
    assert "variabilidade" in resumo
    assert "sem inferência de função" in resumo
    assert "Evolucoes registradas" in resumo

    buffer = gerar_doc_completo(
        "Paciente Teste",
        df_hist,
        df_hist,
        df_alvos,
        pd.DataFrame(),
        df_beh,
        "Media geral dos objetivos",
        date(2026, 7, 1),
        date(2026, 7, 31),
        date(2026, 8, 1),
        resumo,
        "Resumo dos comportamentos.",
    )
    doc = Document(buffer)
    textos = [paragrafo.text for paragrafo in doc.paragraphs]
    assert any("Resumo clinico do desempenho do paciente no periodo" in texto for texto in textos)
    assert any("Mando com frases" in texto for texto in textos)


def test_ausencia_de_registro_nao_e_descrita_como_ausencia_do_comportamento():
    resumo = gerar_resumo_comportamentos_problema_local(
        pd.DataFrame(), date(2026, 7, 1), date(2026, 7, 31)
    )
    assert "ausência de registro não equivale a ausência do comportamento" in resumo


def test_resumo_de_interferentes_nao_inventa_funcao():
    dados = pd.DataFrame(
        {
            "comportamento": ["Fuga", "Fuga", "Fuga"],
            "date": ["2026-07-01", "2026-07-02", "2026-07-03"],
            "count": [3, 2, 1],
            "rate": [3.0, 2.0, 1.0],
        }
    )
    resumo = gerar_resumo_comportamentos_problema_local(
        dados, date(2026, 7, 1), date(2026, 7, 31)
    )
    assert "tendência redução" in resumo
    assert "função do comportamento não pode ser determinada" in resumo
