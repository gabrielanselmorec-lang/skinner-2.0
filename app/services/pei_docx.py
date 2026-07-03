"""Geração de documentos Word (DOCX) do PEI.

Depende de:
- app.services.pei_rules  (funções puras de regra clínica)
- app.web.api_client      (ask_clinical_agent para textos com IA)
"""
import io
import re
from datetime import datetime, timedelta

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.ticker as ticker
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import streamlit as st

from app.services.pei_rules import (
    calcular_evolucao_trimestral,
    classificar_area_pei,
    filtrar_periodo_pei,
    formatar_lista_alvos_corrida,
    limiar_programa_pei,
    limpar_texto_pei,
    resposta_ia_invalida,
    resumo_alvos_programa,
    status_geral_alvos,
    texto_objetivo_programa,
    verificar_alvos_clean,
)

START_DATE = datetime(2026, 3, 27)


# ---------------------------------------------------------------------------
# Helpers de estilo DOCX
# ---------------------------------------------------------------------------

def aplicar_fonte_pei(doc):
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass

    def _formatar_paragrafo(paragraph):
        for run in paragraph.runs:
            if run._element.xpath(".//*[local-name()='drawing']") or run._element.xpath(".//*[local-name()='pict']"):
                continue
            run.text = limpar_texto_pei(run.text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            try:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            except Exception:
                pass

    for paragraph in doc.paragraphs:
        _formatar_paragrafo(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _formatar_paragrafo(paragraph)


def set_cell_text(cell, text):
    cell.text = limpar_texto_pei(text)


def aplicar_estilo_tabela_seguro(table, nomes=("Table Grid", "Tabela Grade", "Grade da Tabela")):
    for nome in nomes:
        try:
            table.style = nome
            return
        except KeyError:
            continue
        except Exception:
            return


def aplicar_bordas_grade_tabela(table, color="000000", size="8"):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = f"w:{edge}"
                element = tc_borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), size)
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), color)


def limpar_linhas_tabela(table):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""


def garantir_colunas_tabela(table, total_colunas):
    while len(table.rows[0].cells) < total_colunas:
        table.add_column(Inches(1.1))


def deixar_celula_negrito(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def aplicar_cor_fundo_celula(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def cor_status_pei(status):
    status_cf = str(status or "").casefold()
    if "atingiu" in status_cf or "atingido" in status_cf:
        return "C6EFCE"
    if "avançou" in status_cf or "avancou" in status_cf:
        return "BDD7EE"
    if "estabilizou" in status_cf or "estagnou" in status_cf:
        return "FFF2CC"
    if "agravou" in status_cf:
        return "F4CCCC"
    return "E7E6E6"


def cor_status_css_pei(status):
    return f"background-color: #{cor_status_pei(status)}"


def estilizar_status_preview_pei(df):
    styler = df.style
    if hasattr(styler, "map"):
        return styler.map(cor_status_css_pei, subset=["Status"])
    return styler.apply(
        lambda coluna: [cor_status_css_pei(valor) for valor in coluna],
        subset=["Status"],
    )


# ---------------------------------------------------------------------------
# Tabelas de objetivos e desempenho
# ---------------------------------------------------------------------------

def distribuir_objetivos_por_area(df_prog, df_beh):
    areas = {1: [], 2: [], 3: [], 4: [], 5: []}
    df_prog_unicos = df_prog.drop_duplicates(subset=["programa"]) if not df_prog.empty else pd.DataFrame()
    for _, row in df_prog_unicos.iterrows():
        programa = str(row.get("programa", "")).strip()
        objetivo = texto_objetivo_programa(row, verificar_alvos_clean)
        area = classificar_area_pei(programa, objetivo)
        areas[area].append({"programa": programa, "texto": objetivo})

    if not df_beh.empty and "comportamento" in df_beh.columns:
        comportamentos = sorted([str(v) for v in df_beh["comportamento"].dropna().unique() if str(v).strip()])
        if comportamentos:
            texto = (
                "Reduzir a taxa dos comportamentos interferentes monitorados "
                f"({', '.join(comportamentos[:8])}) com base nos registros clínicos do período."
            )
            areas[4].insert(0, {"programa": "Comportamentos interferentes", "texto": texto})
    return areas


def preencher_tabela_objetivos(table, area_num, itens):
    garantir_colunas_tabela(table, 2)
    limpar_linhas_tabela(table)

    headers = ["Objetivo", "Descrição"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header)
        deixar_celula_negrito(table.rows[0].cells[idx])

    for item_index, item in enumerate(itens):
        cells = table.add_row().cells
        set_cell_text(cells[0], f"Objetivo {area_num}.{item_index + 1}")
        set_cell_text(cells[1], item.get("texto", ""))
        deixar_celula_negrito(cells[0])

    if not itens:
        cells = table.add_row().cells
        set_cell_text(cells[0], "-")
        set_cell_text(cells[1], "Sem objetivos cadastrados para esta área.")

    aplicar_estilo_tabela_seguro(table)
    aplicar_bordas_grade_tabela(table)


def desempenho_trimestral_para_programa(programa, df_hist, limiar=90, start_date=None):
    if start_date is None:
        start_date = START_DATE
    if df_hist.empty or not programa:
        return None, []
    hist = df_hist[df_hist["programa"] == programa].copy()
    if hist.empty:
        return None, []
    hist["date_pd"] = pd.to_datetime(hist["date"], errors="coerce")
    hist = hist.dropna(subset=["date_pd"]).sort_values("date_pd")
    if hist.empty:
        return None, []
    if "phase" not in hist.columns:
        hist["phase"] = ""
    mask_lb = hist["phase"].astype(str).str.contains(r"linha de base|\blb\b|baseline|sondagem", case=False, na=False)
    linha_de_base = hist[mask_lb]["independent_rate"].mean() if not hist[mask_lb].empty else hist["independent_rate"].iloc[0]

    trimestres = [
        (start_date, start_date + timedelta(days=90), "1. DATA"),
        (start_date + timedelta(days=90), start_date + timedelta(days=180), "2. DATA"),
        (start_date + timedelta(days=180), start_date + timedelta(days=270), "3. DATA"),
        (start_date + timedelta(days=270), start_date + timedelta(days=360), "4. DATA"),
    ]
    desempenho_anterior = linha_de_base
    linhas = []
    for ini, fim, rotulo in trimestres:
        df_tri = hist[(hist["date_pd"] >= pd.Timestamp(ini)) & (hist["date_pd"] < pd.Timestamp(fim))]
        if df_tri.empty:
            linhas.append((rotulo, "-", "Sem avaliações"))
            continue
        media = df_tri["independent_rate"].mean()
        codigo = calcular_evolucao_trimestral(desempenho_anterior, media, limiar)
        linhas.append((rotulo.replace("DATA", fim.strftime("%d/%m/%Y")), f"{media:.1f}%", codigo))
        desempenho_anterior = media
    return linha_de_base, linhas


def preencher_tabela_desempenho(table, area_num, itens, df_hist, df_alvos, df_alvos_periodo, df_lib, periodo_inicio, periodo_fim):
    garantir_colunas_tabela(table, 4)
    limpar_linhas_tabela(table)

    headers = ["Objetivo", "Período avaliado", "Desempenho inicial", "Status"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header)
        deixar_celula_negrito(table.rows[0].cells[idx])

    periodo = f"{periodo_inicio.strftime('%d/%m/%Y')} a {periodo_fim.strftime('%d/%m/%Y')}"
    if not itens:
        row = table.add_row()
        set_cell_text(row.cells[0], "-")
        set_cell_text(row.cells[1], periodo)
        set_cell_text(row.cells[2], "-")
        set_cell_text(row.cells[3], "Sem registro")
        aplicar_cor_fundo_celula(row.cells[3], cor_status_pei("Sem registro"))
        row_alvos = table.add_row()
        set_cell_text(row_alvos.cells[0], "Alvos trabalhados")
        merged = row_alvos.cells[1].merge(row_alvos.cells[3])
        set_cell_text(merged, "Sem alvos trabalhados no período")
    else:
        for item_index, item in enumerate(itens):
            programa = item.get("programa", "")
            limiar = limiar_programa_pei(programa, df_lib)
            linha_de_base, _ = desempenho_trimestral_para_programa(programa, df_hist, limiar)
            desempenho_inicial = f"{linha_de_base:.1f}%" if pd.notna(linha_de_base) else "-"
            codigo = f"Objetivo {area_num}.{item_index + 1}"
            alvos = resumo_alvos_programa(df_alvos, programa, df_lib, df_alvos_periodo)
            status_geral = status_geral_alvos(alvos)

            row = table.add_row()
            set_cell_text(row.cells[0], codigo)
            set_cell_text(row.cells[1], periodo)
            set_cell_text(row.cells[2], desempenho_inicial)
            set_cell_text(row.cells[3], status_geral)
            deixar_celula_negrito(row.cells[0])
            aplicar_cor_fundo_celula(row.cells[3], cor_status_pei(status_geral))

            row_alvos = table.add_row()
            set_cell_text(row_alvos.cells[0], "Alvos trabalhados")
            deixar_celula_negrito(row_alvos.cells[0])
            merged = row_alvos.cells[1].merge(row_alvos.cells[3])
            set_cell_text(merged, formatar_lista_alvos_corrida(alvos))
            aplicar_cor_fundo_celula(merged, cor_status_pei(status_geral))

    aplicar_estilo_tabela_seguro(table)
    aplicar_bordas_grade_tabela(table)


# ---------------------------------------------------------------------------
# Geração de texto clínico (local + IA)
# ---------------------------------------------------------------------------

def coletar_evolucoes_periodo(df_hist_periodo, df_alvos_periodo, df_beh_periodo, limite=30):
    itens = []
    fontes = [
        (df_hist_periodo, "programa", "programa"),
        (df_alvos_periodo, "alvo", "programa"),
        (df_beh_periodo, "interferente", "comportamento"),
    ]
    for df_fonte, tipo, nome_col in fontes:
        if df_fonte.empty or "evolution" not in df_fonte.columns:
            continue
        for _, row in df_fonte.iterrows():
            texto = limpar_texto_pei(row.get("evolution", ""))
            if not texto or texto in {"0", "0.0", "nan", "None"}:
                continue
            nome = limpar_texto_pei(row.get(nome_col, ""))
            alvo = limpar_texto_pei(row.get("target_name", "")) if tipo == "alvo" else ""
            data = row.get("date", "")
            label = f"{tipo}: {nome}"
            if alvo:
                label += f" / {alvo}"
            itens.append(f"- {label}; data={data}; evolucao={texto}")
    return "\n".join(itens[:limite])


def gerar_resumo_trimestral_local(df_hist_periodo, df_alvos_periodo, df_beh_periodo, objetivo_grafico, inicio, fim):
    periodo = f"{inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}"
    partes = [f"Analise do periodo de {periodo}."]

    df_obj = df_hist_periodo.copy()
    if objetivo_grafico != "Media geral dos objetivos" and not df_obj.empty:
        df_obj = df_obj[df_obj["programa"] == objetivo_grafico]

    if not df_obj.empty:
        media_ind = df_obj["independent_rate"].mean()
        programas = df_obj["programa"].nunique() if "programa" in df_obj.columns else 0
        alvo_txt = objetivo_grafico if objetivo_grafico != "Media geral dos objetivos" else f"{programas} objetivos monitorados"
        partes.append(f"No periodo, {alvo_txt} apresentou media de independencia de {media_ind:.1f}%.")
    else:
        partes.append("Nao houve registros de objetivos no ciclo selecionado.")

    if not df_alvos_periodo.empty and "target_name" in df_alvos_periodo.columns:
        total_alvos = df_alvos_periodo["target_name"].nunique()
        media_alvos = df_alvos_periodo["independent_rate"].mean() if "independent_rate" in df_alvos_periodo.columns else None
        if pd.notna(media_alvos):
            partes.append(f"Foram registrados {total_alvos} alvos no periodo, com media de independencia de {media_alvos:.1f}%.")
        else:
            partes.append(f"Foram registrados {total_alvos} alvos no periodo.")

    if not df_beh_periodo.empty and "comportamento" in df_beh_periodo.columns:
        total = df_beh_periodo["count"].sum() if "count" in df_beh_periodo.columns else 0
        taxa = df_beh_periodo["rate"].mean() if "rate" in df_beh_periodo.columns else None
        principais = (
            ", ".join(df_beh_periodo.groupby("comportamento")["count"].sum().sort_values(ascending=False).head(3).index.astype(str))
            if "count" in df_beh_periodo.columns
            else ", ".join(df_beh_periodo["comportamento"].dropna().astype(str).unique()[:3])
        )
        trecho = f"Os comportamentos interferentes somaram {total:.0f} ocorrencias"
        if pd.notna(taxa):
            trecho += f", com taxa media de {taxa:.2f}"
        if principais:
            trecho += f". Principais registros: {principais}"
        partes.append(trecho + ".")
    else:
        partes.append("Nao houve comportamentos interferentes registrados no ciclo selecionado.")

    evolucoes = coletar_evolucoes_periodo(df_hist_periodo, df_alvos_periodo, df_beh_periodo, limite=12)
    if evolucoes:
        partes.append("Evolucoes registradas no periodo:\n" + evolucoes)

    return "\n\n".join(partes)


def gerar_texto_trimestral_pei(
    nome_paciente, df_hist_periodo, df_alvos_periodo, df_beh_periodo,
    objetivo_grafico, inicio, fim, usar_ia=True, ask_agent_fn=None
):
    resumo_local = gerar_resumo_trimestral_local(
        df_hist_periodo, df_alvos_periodo, df_beh_periodo, objetivo_grafico, inicio, fim
    )
    if not usar_ia or ask_agent_fn is None:
        return resumo_local

    pergunta = (
        "Responda exclusivamente em portugues brasileiro, em paragrafo corrido. "
        "Nao use codigo, JSON, markdown, listas tecnicas, nomes de variaveis ou termos em ingles. "
        "Responda apenas com o texto final do PEI, sem repetir estas instrucoes. "
        "Escreva um texto tecnico, objetivo e cauteloso para inserir no PEI. "
        f"O periodo selecionado e de {inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}. "
        "Leia e sintetize as evolucoes/anotacoes registradas pelos terapeutas no periodo, "
        "sem citar nomes de terapeutas. "
        "Explique como foram os objetivos no periodo selecionado, destacando "
        "progresso, estabilidade, agravamento, alvos relevantes e comportamentos interferentes. "
        "Use apenas os dados e evolucoes do periodo e nao invente informacoes. "
        f"Objetivo em destaque para o grafico: {objetivo_grafico}."
    )

    def _parece_prompt_ia(texto):
        texto_cf = str(texto or "").casefold()
        marcadores = [
            "responda apenas com o texto final do pei",
            "escreva um texto tecnico",
            "explique como foram os objetivos",
            "objetivo em destaque para o grafico",
        ]
        return any(m in texto_cf for m in marcadores)

    try:
        resposta = ask_agent_fn(nome_paciente, pergunta, start_date=inicio, end_date=fim)
        modo = str(resposta.get("modo", "")).casefold()
        if modo in {"busca_local", "busca_local_fallback", "limite_ia", "sem_resultados"}:
            return resumo_local
        texto = limpar_texto_pei(resposta.get("resposta", ""))
        if texto and not _parece_prompt_ia(texto) and not resposta_ia_invalida(texto):
            return texto
    except Exception as exc:
        st.warning(f"Nao foi possivel gerar o texto com IA. Usei um resumo automatico local. Detalhe: {exc}")
    return resumo_local


def gerar_resumo_comportamentos_problema_local(df_beh_periodo, inicio, fim):
    periodo = f"{inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}"
    if df_beh_periodo.empty or "comportamento" not in df_beh_periodo.columns:
        return f"No periodo de {periodo}, nao houve registros de comportamentos problema."

    partes = [f"Resumo dos comportamentos problema no periodo de {periodo}."]
    if "count" in df_beh_periodo.columns:
        total_por_comp = df_beh_periodo.groupby("comportamento")["count"].sum().sort_values(ascending=False)
        if not total_por_comp.empty:
            principais = "; ".join([f"{comp}: {valor:.0f} ocorrencias" for comp, valor in total_por_comp.head(5).items()])
            partes.append(f"Maiores registros quantitativos: {principais}.")
    if "rate" in df_beh_periodo.columns:
        taxa_por_comp = df_beh_periodo.groupby("comportamento")["rate"].mean().sort_values(ascending=False)
        if not taxa_por_comp.empty:
            taxas = "; ".join([f"{comp}: taxa media {valor:.2f}" for comp, valor in taxa_por_comp.head(5).items()])
            partes.append(f"Taxas medias observadas: {taxas}.")

    evolucoes = coletar_evolucoes_periodo(pd.DataFrame(), pd.DataFrame(), df_beh_periodo, limite=10)
    if evolucoes:
        partes.append("Evolucoes relevantes registradas no periodo:\n" + evolucoes)
    else:
        partes.append("Nao foram encontradas evolucoes textuais especificas para os comportamentos no periodo.")
    return "\n\n".join(partes)


def gerar_resumo_comportamentos_problema_pei(
    nome_paciente, df_beh_periodo, inicio, fim, usar_ia=True, ask_agent_fn=None
):
    resumo_local = gerar_resumo_comportamentos_problema_local(df_beh_periodo, inicio, fim)
    if not usar_ia or ask_agent_fn is None:
        return resumo_local

    evolucoes = coletar_evolucoes_periodo(pd.DataFrame(), pd.DataFrame(), df_beh_periodo, limite=20)
    pergunta = (
        "Responda exclusivamente em portugues brasileiro, em paragrafo corrido. "
        "Nao use codigo, JSON, markdown, listas tecnicas, nomes de variaveis ou termos em ingles. "
        "Responda apenas com o texto final para o PEI. "
        "Faca um resumo tecnico, objetivo e cauteloso dos comportamentos problema/interferentes "
        "do periodo selecionado. Use os registros de evolucao para filtrar o que de fato ocorreu, "
        "sem inventar informacoes e sem citar nomes de terapeutas. "
        f"Periodo: {inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}. "
        f"Resumo automatico dos dados: {resumo_local[:1800]} "
        f"Evolucoes do periodo: {evolucoes[:1500]}"
    )
    try:
        resposta = ask_agent_fn(nome_paciente, pergunta, start_date=inicio, end_date=fim)
        modo = str(resposta.get("modo", "")).casefold()
        if modo in {"busca_local", "busca_local_fallback", "limite_ia", "sem_resultados"}:
            return resumo_local
        texto = limpar_texto_pei(resposta.get("resposta", ""))
        if texto and not resposta_ia_invalida(texto):
            return texto
    except Exception as exc:
        st.warning(f"Nao foi possivel gerar o resumo dos comportamentos com IA. Usei um resumo local. Detalhe: {exc}")
    return resumo_local


# ---------------------------------------------------------------------------
# Funções de inserção de seções no DOCX
# ---------------------------------------------------------------------------

def adicionar_texto_trimestral(doc, texto, inicio, fim):
    doc.add_page_break()
    doc.add_heading("Analise dos objetivos no periodo", level=1)
    p_periodo = doc.add_paragraph()
    p_periodo.add_run("Periodo: ").bold = True
    p_periodo.add_run(f"{inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}")

    texto_limpo = limpar_texto_pei(texto)
    if not texto_limpo:
        texto_limpo = "Espaco reservado para descricao clinica do desempenho dos objetivos no periodo."
    for bloco in re.split(r"\n\s*\n", texto_limpo):
        if bloco.strip():
            doc.add_paragraph(bloco.strip())


def adicionar_resumo_clinico_periodo(doc, texto, inicio, fim):
    doc.add_page_break()
    doc.add_heading("Resumo clínico do período observado", level=1)
    p_periodo = doc.add_paragraph()
    p_periodo.add_run("Periodo observado: ").bold = True
    p_periodo.add_run(f"{inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}")

    texto_limpo = limpar_texto_pei(texto)
    if not texto_limpo:
        texto_limpo = "Sem registros suficientes para sintetizar o período observado."
    for bloco in re.split(r"\n\s*\n", texto_limpo):
        if bloco.strip():
            doc.add_paragraph(bloco.strip())


def adicionar_grafico_objetivo_periodo(doc, df_hist, objetivo_grafico, inicio, fim_exclusivo):
    doc.add_heading("Graficos dos objetivos", level=2)
    df_periodo = filtrar_periodo_pei(df_hist, inicio, fim_exclusivo)

    if df_periodo.empty:
        doc.add_paragraph("Sem registros de objetivos no periodo selecionado.")
        return

    df_obj = df_periodo.copy()
    titulo = "Media geral dos objetivos"
    if objetivo_grafico != "Media geral dos objetivos":
        df_selecionado = df_periodo[df_periodo["programa"] == objetivo_grafico].copy()
        if not df_selecionado.empty:
            df_obj = df_selecionado
            titulo = objetivo_grafico
        else:
            doc.add_paragraph(
                "O objetivo selecionado nao teve registros no periodo; abaixo seguem os objetivos com dados no periodo."
            )

    df_plot = (
        df_obj.groupby("date_pd", as_index=False)
        .agg({"independent_rate": "mean"})
        .sort_values("date_pd")
    )
    if df_plot.empty:
        doc.add_paragraph("Sem dados numericos suficientes para gerar o grafico do objetivo.")
        return

    plt.figure(figsize=(7, 4))
    plt.plot(df_plot["date_pd"], df_plot["independent_rate"], marker="o", label="Independencia")
    plt.title(limpar_texto_pei(titulo)[:90])
    plt.ylabel("Percentual")
    plt.ylim(0, 100)
    ax = plt.gca()
    ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
    plt.gcf().autofmt_xdate(rotation=45)
    plt.legend()
    plt.tight_layout()
    buf_obj = io.BytesIO()
    plt.savefig(buf_obj, format="png")
    plt.close()
    buf_obj.seek(0)
    doc.add_picture(buf_obj, width=Inches(5.7))

    df_barras = (
        df_periodo.groupby("programa", as_index=False)["independent_rate"]
        .mean()
        .sort_values("independent_rate", ascending=True)
        .tail(15)
    )
    if not df_barras.empty:
        plt.figure(figsize=(7, 5))
        plt.barh(df_barras["programa"], df_barras["independent_rate"], color="#2E86C1")
        plt.title("Media de independencia por objetivo no periodo")
        plt.xlabel("Independencia media (%)")
        plt.xlim(0, 100)
        plt.tight_layout()
        buf_bar = io.BytesIO()
        plt.savefig(buf_bar, format="png", dpi=160)
        plt.close()
        buf_bar.seek(0)
        doc.add_picture(buf_bar, width=Inches(5.8))


def adicionar_graficos_interferentes_periodo(doc, df_beh, inicio, fim_exclusivo):
    doc.add_heading("Graficos de comportamentos interferentes", level=2)
    df_beh_periodo = filtrar_periodo_pei(df_beh, inicio, fim_exclusivo)
    if df_beh_periodo.empty:
        doc.add_paragraph("Sem comportamentos interferentes registrados no periodo selecionado.")
        return

    if "count" in df_beh_periodo.columns:
        df_totais = df_beh_periodo.groupby("comportamento")["count"].sum().reset_index()
        if not df_totais.empty and df_totais["count"].sum() > 0:
            plt.figure(figsize=(6.5, 4))
            plt.bar(df_totais["comportamento"], df_totais["count"], color="#E74C3C")
            plt.title("Frequencia total no periodo")
            plt.xticks(rotation=45, ha="right")
            plt.tight_layout()
            buf_tot = io.BytesIO()
            plt.savefig(buf_tot, format="png")
            plt.close()
            buf_tot.seek(0)
            doc.add_picture(buf_tot, width=Inches(5.3))

    if "rate" in df_beh_periodo.columns:
        df_taxa = df_beh_periodo.groupby("comportamento")["rate"].mean().reset_index()
        if not df_taxa.empty and df_taxa["rate"].sum() > 0:
            plt.figure(figsize=(6.5, 4))
            plt.bar(df_taxa["comportamento"], df_taxa["rate"], color="#F39C12")
            plt.title("Taxa media no periodo")
            plt.xticks(rotation=45, ha="right")
            plt.tight_layout()
            buf_taxa = io.BytesIO()
            plt.savefig(buf_taxa, format="png")
            plt.close()
            buf_taxa.seek(0)
            doc.add_picture(buf_taxa, width=Inches(5.3))

    plt.figure(figsize=(7, 4))
    for comportamento in df_beh_periodo["comportamento"].dropna().unique():
        df_c = df_beh_periodo[df_beh_periodo["comportamento"] == comportamento].sort_values("date_pd")
        y_col = "rate" if "rate" in df_c.columns else "count"
        plt.plot(df_c["date_pd"], df_c[y_col], marker="o", label=str(comportamento))
    plt.title("Evolucao dos interferentes no periodo")
    ax = plt.gca()
    ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
    plt.gcf().autofmt_xdate(rotation=45)
    plt.legend(bbox_to_anchor=(1.05, 1), loc="upper left")
    plt.tight_layout()
    buf_linha = io.BytesIO()
    plt.savefig(buf_linha, format="png")
    plt.close()
    buf_linha.seek(0)
    doc.add_picture(buf_linha, width=Inches(5.7))


def adicionar_graficos_modelo(doc, df_hist, df_alvos, df_beh, objetivo_grafico, periodo_inicio, periodo_fim_exclusivo):
    doc.add_page_break()
    doc.add_heading("ANEXOS GRÁFICOS", level=1)
    adicionar_grafico_objetivo_periodo(doc, df_hist, objetivo_grafico, periodo_inicio, periodo_fim_exclusivo)
    adicionar_graficos_interferentes_periodo(doc, df_beh, periodo_inicio, periodo_fim_exclusivo)

    if not df_hist.empty:
        doc.add_heading("Resumo geral dos programas", level=2)
        df_graf = df_hist.copy()
        df_graf["date_pd"] = pd.to_datetime(df_graf["date"], errors="coerce")
        df_graf = df_graf.dropna(subset=["date_pd"])
        if not df_graf.empty:
            df_linha = (
                df_graf.groupby("date_pd", as_index=False)
                .agg({"independent_rate": "mean", "prompt_rate": "mean"})
                .sort_values("date_pd")
            )
            if not df_linha.empty:
                plt.figure(figsize=(7, 4))
                plt.plot(df_linha["date_pd"], df_linha["independent_rate"], marker="o", label="Independência média")
                plt.plot(df_linha["date_pd"], df_linha["prompt_rate"], marker="o", label="Ajuda média")
                plt.title("Evolução média dos programas")
                plt.ylabel("Percentual")
                plt.ylim(0, 100)
                ax = plt.gca()
                ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
                ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
                plt.gcf().autofmt_xdate(rotation=45)
                plt.legend()
                plt.tight_layout()
                buf = io.BytesIO()
                plt.savefig(buf, format="png")
                plt.close()
                buf.seek(0)
                doc.add_picture(buf, width=Inches(5.7))

    if not df_beh.empty:
        df_beh_g = df_beh.copy()
        df_beh_g["date_pd"] = pd.to_datetime(df_beh_g["date"], errors="coerce")
        df_beh_g = df_beh_g.dropna(subset=["date_pd"])
        if not df_beh_g.empty:
            doc.add_heading("Comportamentos interferentes", level=2)
            df_totais = df_beh_g.groupby("comportamento")["count"].sum().reset_index()
            if not df_totais.empty:
                plt.figure(figsize=(6, 4))
                plt.bar(df_totais["comportamento"], df_totais["count"], color="#E74C3C")
                plt.title("Frequência total de ocorrências")
                plt.xticks(rotation=45, ha="right")
                plt.tight_layout()
                buf_tot = io.BytesIO()
                plt.savefig(buf_tot, format="png")
                plt.close()
                buf_tot.seek(0)
                doc.add_picture(buf_tot, width=Inches(5.2))

            if "rate" in df_beh_g.columns:
                df_taxa = df_beh_g.groupby("comportamento")["rate"].mean().reset_index()
                if not df_taxa.empty:
                    plt.figure(figsize=(6, 4))
                    plt.bar(df_taxa["comportamento"], df_taxa["rate"], color="#F39C12")
                    plt.title("Taxa média")
                    plt.xticks(rotation=45, ha="right")
                    plt.tight_layout()
                    buf_taxa = io.BytesIO()
                    plt.savefig(buf_taxa, format="png")
                    plt.close()
                    buf_taxa.seek(0)
                    doc.add_picture(buf_taxa, width=Inches(5.2))

            plt.figure(figsize=(7, 4))
            for comportamento in df_beh_g["comportamento"].dropna().unique():
                df_c = df_beh_g[df_beh_g["comportamento"] == comportamento].sort_values("date_pd")
                plt.plot(df_c["date_pd"], df_c["count"], marker="o", label=comportamento)
            plt.title("Evolução ao longo do tempo")
            ax = plt.gca()
            ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
            ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
            plt.gcf().autofmt_xdate(rotation=45)
            plt.legend(bbox_to_anchor=(1.05, 1), loc="upper left")
            plt.tight_layout()
            buf_linha = io.BytesIO()
            plt.savefig(buf_linha, format="png")
            plt.close()
            buf_linha.seek(0)
            doc.add_picture(buf_linha, width=Inches(5.7))


# ---------------------------------------------------------------------------
# Funções públicas principais
# ---------------------------------------------------------------------------

def gerar_doc_modelo_pei(
    nome_paciente,
    df_prog,
    df_hist,
    df_alvos,
    df_lib,
    df_beh,
    objetivo_grafico,
    periodo_inicio,
    periodo_fim,
    periodo_fim_exclusivo,
    texto_analise_trimestral,
    texto_resumo_comportamentos,
    pei_template_path="",
    start_date=None,
):
    if start_date is None:
        start_date = START_DATE
    doc = Document(pei_template_path)
    aplicar_fonte_pei(doc)
    areas = distribuir_objetivos_por_area(df_prog, df_beh)
    df_alvos_periodo = filtrar_periodo_pei(df_alvos, periodo_inicio, periodo_fim_exclusivo)

    if doc.tables:
        dados = (
            "DADOS DE IDENTIFICAÇÃO\n"
            f"Nome do(a) Aprendiz: {nome_paciente}\n"
            "Data de Nascimento: N/I | Idade: N/I\n"
            f"Data de Elaboração do Plano: {datetime.now().strftime('%m/%Y')}\n"
            "Equipe Responsável: Sistema Skinner\n"
            "Equipe Atual: Aplicável em caso de mudança de equipe\n"
            "Instrumento(s) de Avaliação Utilizado(s): bHave, ABLLS-R e registros clínicos\n"
            f"Data da Avaliação: {start_date.strftime('%d/%m/%Y')}\n"
            f"Data da Finalização: {datetime.now().strftime('%d/%m/%Y')}"
        )
        set_cell_text(doc.tables[0].cell(1, 0), dados)

    for area_num in range(1, 6):
        table_index = area_num
        if table_index < len(doc.tables):
            preencher_tabela_objetivos(
                doc.tables[table_index],
                area_num,
                areas.get(area_num, []),
            )

    desempenho_tables = {1: 6, 2: 7, 3: 8}
    for area_num, table_index in desempenho_tables.items():
        if table_index < len(doc.tables):
            preencher_tabela_desempenho(
                doc.tables[table_index],
                area_num,
                areas.get(area_num, []),
                df_hist,
                df_alvos,
                df_alvos_periodo,
                df_lib,
                periodo_inicio,
                periodo_fim,
            )

    adicionar_resumo_clinico_periodo(doc, texto_resumo_comportamentos, periodo_inicio, periodo_fim)
    adicionar_graficos_modelo(doc, df_hist, df_alvos, df_beh, objetivo_grafico, periodo_inicio, periodo_fim_exclusivo)
    aplicar_fonte_pei(doc)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def gerar_doc_completo(
    nome_paciente,
    df_prog,
    df_hist,
    df_alvos,
    df_lib,
    df_beh,
    objetivo_grafico,
    periodo_inicio,
    periodo_fim,
    periodo_fim_exclusivo,
    texto_analise_trimestral,
    texto_resumo_comportamentos,
    pei_template_path="",
    start_date=None,
):
    import os
    if start_date is None:
        start_date = START_DATE
    if pei_template_path and os.path.exists(pei_template_path):
        return gerar_doc_modelo_pei(
            nome_paciente, df_prog, df_hist, df_alvos, df_lib, df_beh,
            objetivo_grafico, periodo_inicio, periodo_fim, periodo_fim_exclusivo,
            texto_analise_trimestral, texto_resumo_comportamentos,
            pei_template_path=pei_template_path, start_date=start_date,
        )

    from app.services.pei_rules import (
        calcular_evolucao_trimestral as _calc_evo,
        filtrar_periodo_pei as _filtrar,
        verificar_alvos_clean as _verificar,
    )

    doc = Document()
    aplicar_fonte_pei(doc)

    status_objetivos = {"ATG": 0, "AVA": 0, "EST": 0, "AGR": 0, "-": 0}
    total_alvos_trabalhados = 0
    total_alvos_atingidos = 0

    title = doc.add_heading("PLANO DE ENSINO INDIVIDUALIZADO (PEI)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("DADOS DE IDENTIFICAÇÃO", level=1)
    t_id = doc.add_table(rows=4, cols=2)
    aplicar_estilo_tabela_seguro(t_id)
    t_id.rows[0].cells[0].text = f"Nome do(a) Aprendiz: {nome_paciente}"
    t_id.rows[1].cells[0].text = f"Data de Início da Análise: {start_date.strftime('%d/%m/%Y')}"
    t_id.rows[2].cells[0].text = f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y')}"
    t_id.rows[3].cells[0].text = "Equipe Responsável: Sistema Skinner"

    df_alvos_periodo = _filtrar(df_alvos, periodo_inicio, periodo_fim_exclusivo)
    adicionar_resumo_clinico_periodo(doc, texto_resumo_comportamentos, periodo_inicio, periodo_fim)

    doc.add_heading("OBJETIVOS DE INTERVENÇÃO", level=1)
    doc.add_paragraph("(As metas devem ser Específicas, Mensuráveis, Alcançáveis, Relevantes e com Prazo definido.)\n")

    df_prog_periodo = _filtrar(df_prog, periodo_inicio, periodo_fim_exclusivo)
    df_prog_unicos = df_prog_periodo.drop_duplicates(subset=["programa"])
    areas = distribuir_objetivos_por_area(df_prog_periodo, df_beh)

    for area_num in range(1, 6):
        doc.add_heading(f"Area {area_num}", level=2)
        tabela_area = doc.add_table(rows=1, cols=2)
        preencher_tabela_objetivos(tabela_area, area_num, areas.get(area_num, []))

    doc.add_heading("ALVOS TRABALHADOS E DESEMPENHO", level=1)
    for area_num in range(1, 6):
        doc.add_heading(f"Area {area_num}", level=2)
        tabela_alvos = doc.add_table(rows=18, cols=4)
        preencher_tabela_desempenho(
            tabela_alvos, area_num, areas.get(area_num, []),
            df_hist, df_alvos, df_alvos_periodo, df_lib, periodo_inicio, periodo_fim,
        )

    doc.add_heading("DESEMPENHO POR ÁREA", level=1)
    trimestres = [
        (start_date, start_date + timedelta(days=90), "Trimestre 1"),
        (start_date + timedelta(days=90), start_date + timedelta(days=180), "Trimestre 2"),
        (start_date + timedelta(days=180), start_date + timedelta(days=270), "Trimestre 3"),
        (start_date + timedelta(days=270), start_date + timedelta(days=360), "Trimestre 4"),
    ]

    for _, row in df_prog_unicos.iterrows():
        prog_nome = row["programa"]
        limiar = row.get("mastery_threshold_percent", 90)
        if pd.isna(limiar) or not limiar:
            limiar = 90

        doc.add_heading(f"{prog_nome}", level=3)

        hist_p = df_hist[df_hist["programa"] == prog_nome].copy()
        linha_de_base = None
        media_int_geral = None
        status_lb = ""

        if not hist_p.empty:
            hist_p["date_pd"] = pd.to_datetime(hist_p["date"])
            if "phase" not in hist_p.columns:
                hist_p["phase"] = ""
            mask_hist_lb = hist_p["phase"].astype(str).str.contains(
                r"linha de base|\blb\b|baseline|sondagem", case=False, na=False
            )
            fases_lb = hist_p[mask_hist_lb]
            fases_int = hist_p[~mask_hist_lb]

            if not fases_lb.empty:
                linha_de_base = fases_lb["independent_rate"].mean()
                if linha_de_base >= limiar:
                    status_lb = "Objetivo atingido ja na Linha de Base"
                elif not fases_int.empty:
                    status_lb = "Em Intervencao (avancou apos Linha de Base)"
                else:
                    status_lb = "Avaliacao em andamento (em Linha de Base)"
            else:
                linha_de_base = hist_p.iloc[0]["independent_rate"]
                status_lb = "Em Intervencao (sem registro formal de LB no periodo)"

            if not fases_int.empty:
                media_int_geral = fases_int["independent_rate"].mean()

        p_lb = doc.add_paragraph()
        p_lb.add_run("Média Linha de Base: ").bold = True
        p_lb.add_run(f"{linha_de_base:.1f}%" if pd.notnull(linha_de_base) else "N/A")

        p_int = doc.add_paragraph()
        p_int.add_run("Média Intervenção: ").bold = True
        p_int.add_run(f"{media_int_geral:.1f}%" if pd.notnull(media_int_geral) else "N/A")

        if status_lb:
            p_st = doc.add_paragraph(status_lb)
            p_st.runs[0].italic = True

        table = doc.add_table(rows=1, cols=3)
        aplicar_estilo_tabela_seguro(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Avaliação - Período"
        hdr[1].text = "Desempenho"
        hdr[2].text = "Código de evolução"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True

        desempenho_anterior = linha_de_base
        evo_final_obj = "-"

        for ini, fim, rotulo in trimestres:
            row_cells = table.add_row().cells
            if not hist_p.empty:
                df_tri = hist_p[(hist_p["date_pd"] >= pd.Timestamp(ini)) & (hist_p["date_pd"] < pd.Timestamp(fim))]
                if not df_tri.empty:
                    media = df_tri["independent_rate"].mean()
                    cod_evo = _calc_evo(desempenho_anterior, media, limiar)
                    row_cells[0].text = rotulo
                    row_cells[1].text = f"{media:.1f}%"
                    row_cells[2].text = cod_evo
                    desempenho_anterior = media
                    evo_final_obj = cod_evo.split("(")[1].replace(")", "") if "(" in cod_evo else cod_evo
                    continue
            row_cells[0].text = rotulo
            row_cells[1].text = "-"
            row_cells[2].text = "Sem avaliações"

        if evo_final_obj in status_objetivos:
            status_objetivos[evo_final_obj] += 1
        else:
            status_objetivos["-"] += 1

        if not df_alvos_periodo.empty and "programa" in df_alvos_periodo.columns:
            alvos_prog = df_alvos_periodo[df_alvos_periodo["programa"] == prog_nome].copy()
            p_alvos = doc.add_paragraph()
            p_alvos.add_run("\nAlvos que atingiram a meta no periodo:").bold = True

            if not alvos_prog.empty:
                if "phase" not in alvos_prog.columns:
                    alvos_prog["phase"] = ""
                total_alvos_trabalhados += alvos_prog["target_name"].nunique()
                lb_mastered = alvos_prog[
                    alvos_prog["phase"].astype(str).str.contains(
                        r"linha de base|\blb\b|baseline|sondagem", case=False, na=False
                    ) & (alvos_prog["independent_rate"] >= 100)
                ]["target_name"].unique()
                mask = ~alvos_prog["target_name"].isin(lb_mastered)

                if not alvos_prog[mask].empty:
                    max_rates = alvos_prog[mask].groupby("target_name")["independent_rate"].max()
                    alvos_meta = [(alvo, taxa) for alvo, taxa in max_rates.items() if taxa >= limiar]
                    total_alvos_atingidos += len(alvos_meta)
                    if alvos_meta:
                        for alvo, taxa in sorted(alvos_meta):
                            doc.add_paragraph(f"  - {alvo}: {taxa:.1f}% de independencia")
                    else:
                        doc.add_paragraph("  - Nenhum novo alvo atingiu o criterio de dominio.")
                else:
                    doc.add_paragraph("  - Nenhum novo alvo atingiu o criterio de dominio.")
            else:
                doc.add_paragraph("  - Sem registros granulares de alvos no periodo.")
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("RESUMO CLÍNICO DE EVOLUÇÃO", level=1)
    adicionar_grafico_objetivo_periodo(doc, df_hist, objetivo_grafico, periodo_inicio, periodo_fim_exclusivo)
    adicionar_graficos_interferentes_periodo(doc, df_beh, periodo_inicio, periodo_fim_exclusivo)

    labels_obj, sizes_obj, colors_obj = [], [], []
    color_map = {"ATG": "#2ECC71", "AVA": "#3498DB", "EST": "#F1C40F", "AGR": "#E74C3C", "-": "#95A5A6"}
    nome_map = {"ATG": "Atingiu", "AVA": "Avançou", "EST": "Estabilizou", "AGR": "Agravou", "-": "S/ Dados"}

    for k, v in status_objetivos.items():
        if v > 0:
            labels_obj.append(nome_map.get(k, k))
            sizes_obj.append(v)
            colors_obj.append(color_map.get(k, "#333333"))

    if sizes_obj:
        plt.figure(figsize=(5, 4))
        plt.pie(sizes_obj, labels=labels_obj, colors=colors_obj, autopct="%1.1f%%", startangle=140)
        plt.title("Status Final dos Programas")
        plt.tight_layout()
        buf_obj = io.BytesIO()
        plt.savefig(buf_obj, format="png")
        plt.close()
        buf_obj.seek(0)
        doc.add_picture(buf_obj, width=Inches(4.5))

    if total_alvos_trabalhados > 0:
        plt.figure(figsize=(5, 4))
        bars = plt.bar(
            ["Trabalhados", "Atingiram Meta"],
            [total_alvos_trabalhados, total_alvos_atingidos],
            color=["#3498DB", "#2ECC71"],
        )
        plt.title("Aquisição de Novos Alvos")
        plt.ylabel("Quantidade de Alvos")
        for bar in bars:
            yval = bar.get_height()
            plt.text(bar.get_x() + bar.get_width() / 2, yval + 0.5, int(yval), ha="center", va="bottom", fontweight="bold")
        plt.tight_layout()
        buf_alvos = io.BytesIO()
        plt.savefig(buf_alvos, format="png")
        plt.close()
        buf_alvos.seek(0)
        doc.add_picture(buf_alvos, width=Inches(4.5))

    if not df_hist.empty:
        df_graf = df_hist.copy()
        df_graf["date_pd"] = pd.to_datetime(df_graf["date"], errors="coerce")
        df_graf = df_graf.dropna(subset=["date_pd"])
        if not df_graf.empty:
            doc.add_heading("Graficos de desempenho dos programas", level=2)
            df_linha = (
                df_graf.groupby("date_pd", as_index=False)
                .agg({"independent_rate": "mean", "prompt_rate": "mean"})
                .sort_values("date_pd")
            )
            if not df_linha.empty:
                plt.figure(figsize=(7, 4))
                plt.plot(df_linha["date_pd"], df_linha["independent_rate"], marker="o", label="Independencia media")
                plt.plot(df_linha["date_pd"], df_linha["prompt_rate"], marker="o", label="Ajuda media")
                plt.title("Evolucao media dos programas")
                plt.ylabel("Percentual")
                plt.ylim(0, 100)
                ax = plt.gca()
                ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
                ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
                plt.gcf().autofmt_xdate(rotation=45)
                plt.legend()
                plt.tight_layout()
                buf_prog_linha = io.BytesIO()
                plt.savefig(buf_prog_linha, format="png")
                plt.close()
                buf_prog_linha.seek(0)
                doc.add_picture(buf_prog_linha, width=Inches(5.7))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
