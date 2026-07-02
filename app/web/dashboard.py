import os
import sys
import io
import re
import urllib.parse
from datetime import datetime, date, timedelta

import streamlit as st
import pandas as pd
import requests
import plotly.express as px

# Proteção para usar Matplotlib dentro do Streamlit sem dar conflito de Thread
import matplotlib
matplotlib.use('Agg') 
import matplotlib.pyplot as plt
import matplotlib.dates as mdates   
import matplotlib.ticker as ticker

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ajuste de Caminho para importar o main.py
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../..')))
from main import sincronizar_bhave_api

sys.path.append(os.path.abspath(os.path.dirname(__file__)))
from ablls_view import render_ablls_module

# Configuração da Página
st.set_page_config(page_title="Skinner Project", layout="wide")

# CSS para visual profissional
st.markdown("""
    <style>
    [data-testid="stStatusWidget"] { visibility: hidden; }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    </style>
""", unsafe_allow_html=True)

API_URL = "http://127.0.0.1:8000"
API_TIMEOUT_SECONDS = 90
AGENT_TIMEOUT_SECONDS = int(os.getenv("SKINNER_AGENT_TIMEOUT_SECONDS", "300"))
PEI_TEMPLATE_PATH = os.getenv(
    "SKINNER_PEI_TEMPLATE_PATH",
    r"C:\Users\coord\Desktop\PEI_2026_TEMPLATE_V3.docx",
)

@st.cache_data(ttl=600, show_spinner=False)
def load_data_from_api(paciente):
    try:
        resposta = requests.get(
            f"{API_URL}/api/dados_paciente",
            params={"paciente": paciente},
            timeout=API_TIMEOUT_SECONDS,
        )
        if resposta.status_code != 200: return pd.DataFrame(), pd.DataFrame()
        dados = resposta.json()
        df_p = pd.DataFrame(dados.get("programas", []))
        df_b = pd.DataFrame(dados.get("interferentes", []))
        
        # FORÇA CONVERSÃO NUMÉRICA PARA EVITAR ERRO DE "N/A" NO CÁLCULO DA MÉDIA
        if not df_p.empty:
            if 'date' in df_p.columns: df_p['date'] = pd.to_datetime(df_p['date'], errors='coerce', dayfirst=True).dt.date
            if 'independent_rate' in df_p.columns: df_p['independent_rate'] = pd.to_numeric(df_p['independent_rate'], errors='coerce')
            if 'prompt_rate' in df_p.columns: df_p['prompt_rate'] = pd.to_numeric(df_p['prompt_rate'], errors='coerce')
            
        if not df_b.empty:
            if 'date' in df_b.columns: df_b['date'] = pd.to_datetime(df_b['date'], errors='coerce', dayfirst=True).dt.date
            if 'rate' in df_b.columns: df_b['rate'] = pd.to_numeric(df_b['rate'], errors='coerce')
            if 'count' in df_b.columns: df_b['count'] = pd.to_numeric(df_b['count'], errors='coerce')
            
        return df_p, df_b
    except: return pd.DataFrame(), pd.DataFrame()

@st.cache_data(show_spinner=False)
def load_targets_from_api(paciente, programa):
    try:
        resposta = requests.get(
            f"{API_URL}/api/alvos",
            params={"paciente": paciente, "programa": programa},
            timeout=API_TIMEOUT_SECONDS,
        )
        if resposta.status_code != 200: return pd.DataFrame()
        df = pd.DataFrame(resposta.json())
        
        if not df.empty:
            if 'date' in df.columns: df['date'] = pd.to_datetime(df['date'], errors='coerce', dayfirst=True).dt.date
            if 'independent_rate' in df.columns: df['independent_rate'] = pd.to_numeric(df['independent_rate'], errors='coerce')
            if 'prompt_rate' in df.columns: df['prompt_rate'] = pd.to_numeric(df['prompt_rate'], errors='coerce')
            
        return df
    except: return pd.DataFrame()

@st.cache_data(ttl=600, show_spinner=False)
def load_library_from_api():
    try:
        resposta = requests.get(f"{API_URL}/api/biblioteca", timeout=API_TIMEOUT_SECONDS)
        if resposta.status_code == 200: return pd.DataFrame(resposta.json())
    except: pass
    return pd.DataFrame()

def ask_clinical_agent(paciente, pergunta, start_date=None, end_date=None):
    payload = {
        "paciente": paciente,
        "pergunta": pergunta,
    }
    if start_date:
        payload["data_inicio"] = start_date.isoformat()
    if end_date:
        payload["data_fim"] = end_date.isoformat()
    resposta = requests.post(
        f"{API_URL}/api/agente/clinico",
        json=payload,
        timeout=AGENT_TIMEOUT_SECONDS,
    )
    resposta.raise_for_status()
    return resposta.json()

def render_agent_references(fontes):
    if not fontes:
        return
    st.markdown("#### Referencias bibliograficas")
    for idx, fonte in enumerate(fontes, start=1):
        arquivo = str(fonte.get("fonte") or "Fonte nao informada").strip()
        titulo = str(fonte.get("titulo") or "").strip()
        indice = fonte.get("indice", "")
        partes = [f"arquivo: {arquivo}"]
        if titulo:
            partes.append(f"titulo/capitulo: {titulo}")
        if indice != "":
            partes.append(f"trecho: {indice}")
        st.markdown(f"- **Fonte {idx}:** " + "; ".join(partes) + ".")

st.title("Skinner Project - Análise Clínica")
st.markdown("---")

# Verificação da API
try:
    lista_pacientes = requests.get(f"{API_URL}/api/pacientes", timeout=API_TIMEOUT_SECONDS).json()
    lista_pacientes = sorted(lista_pacientes) 
except:
    st.error("ERRO: A API não está respondendo. Ligue o Uvicorn no terminal.")
    st.stop()

if not lista_pacientes:
    st.warning("Banco de dados vazio.")
    if st.button("Sincronizar bHave (Primeira Vez)"):
        sincronizar_bhave_api()
        st.rerun()
else:
    paciente_sel = st.sidebar.selectbox("Selecione o Paciente", lista_pacientes)
    
    st.sidebar.markdown("---")
    if st.sidebar.button("Sincronizar bHave Agora"):
        with st.spinner("Buscando novos relatórios..."):
            sincronizar_bhave_api()
            st.cache_data.clear() 
            st.rerun()
    
    if paciente_sel:
        df_p_raw, df_b_raw = load_data_from_api(paciente_sel)
        df_lib = load_library_from_api()
        
        if not df_p_raw.empty:
            df_p_raw = df_p_raw[df_p_raw['programa'].str.strip() != ""]
            
            tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "Habilidades", "Decisão", "Interferentes", "PEI", "Avaliações", "Assistente"
    ])

            # ==========================================
            # --- ABA 1: HABILIDADES ---
            # ==========================================
            with tab1:
                st.markdown(f"###  Gestão de Habilidades: **{paciente_sel}**")
                
                if not df_p_raw.empty and not df_p_raw['date'].dropna().empty:
                    min_d_p, max_d_p = df_p_raw['date'].min(), df_p_raw['date'].max()
                else:
                    min_d_p, max_d_p = datetime.today().date(), datetime.today().date()
                
                c_data_p, _ = st.columns([1, 2])
                with c_data_p:
                    periodo_prog = st.date_input(" Período", value=(min_d_p, max_d_p), format="DD/MM/YYYY", key="hab_periodo")
                
                start_p, end_p = (periodo_prog[0], periodo_prog[1]) if len(periodo_prog) == 2 else (min_d_p, max_d_p)
                df_prog = df_p_raw[(df_p_raw['date'] >= start_p) & (df_p_raw['date'] <= end_p)] if not df_p_raw.empty else pd.DataFrame()
                
                if df_prog.empty:
                    st.info("Nenhum dado registrado para este período.")
                else:
                    if 'phase' not in df_prog.columns: df_prog['phase'] = ""
                    # Expressão regular blindada (\blb\b) para não confundir nomes
                    mask_lb = df_prog['phase'].astype(str).str.contains(r'linha de base|\blb\b|baseline|sondagem', case=False, na=False)
                    
                    df_prog['date_pd'] = pd.to_datetime(df_prog['date'])
                    ultima_data_geral = df_prog.groupby('programa')['date_pd'].max()
                    data_corte = pd.to_datetime(end_p) - pd.Timedelta(days=30)
                    programas_ativos = ultima_data_geral[ultima_data_geral >= data_corte].index.tolist()
                    programas_arquivados = ultima_data_geral[ultima_data_geral < data_corte].index.tolist()
                    
                    mostrar_arquivados = st.toggle(" Mostrar Programas Inativos (> 30 dias)")
                    programas_no_periodo = df_prog['programa'].unique().tolist()
                    prog_lista = sorted([p for p in programas_no_periodo if p in (programas_arquivados if mostrar_arquivados else programas_ativos)])
                    
                    if prog_lista:
                        prog_sel = st.selectbox(" 1. Selecione o Programa", ["RESUMO GERAL"] + prog_lista)
                    else:
                        prog_sel = None
                    
                    if prog_sel == "RESUMO GERAL":
                        st.subheader(" Desempenho (Intervenção vs LB)")
                        df_prog_int = df_prog[~mask_lb] # Apenas Intervenção
                        
                        c1, c2, c3 = st.columns(3)
                        media_int = df_prog_int['independent_rate'].mean() if not df_prog_int.empty else None
                        media_lb = df_prog[mask_lb]['independent_rate'].mean() if not df_prog[mask_lb].empty else None
                        
                        c1.metric(" Média TOTAL (Intervenção)", f"{media_int:.1f}%" if pd.notnull(media_int) else "N/A")
                        c2.metric(" Média TOTAL (Linha de Base)", f"{media_lb:.1f}%" if pd.notnull(media_lb) else "N/A")
                        c3.metric(" Programas Ativos", len(programas_ativos))
                        
                        st.markdown("---")
                        if not df_prog_int.empty:
                            df_resumo = df_prog_int[df_prog_int['programa'].isin(prog_lista)].groupby('programa')['independent_rate'].mean().round(1).reset_index().sort_values('independent_rate')
                            fig_resumo = px.bar(df_resumo, y='programa', x='independent_rate', orientation='h', text='independent_rate', color_discrete_sequence=['#2ECC71'], labels={'programa': 'Programa', 'independent_rate': 'Independência (%)'})
                            fig_resumo.update_traces(textposition='outside')
                            fig_resumo.update_layout(xaxis_range=[0, 100])
                            st.plotly_chart(fig_resumo, width='stretch')
                        
                    elif prog_sel:
                        df_p_view = df_prog[df_prog['programa'] == prog_sel].sort_values('date').copy()
                        df_alvos_brutos = load_targets_from_api(paciente_sel, prog_sel)
                        lista_alvos = sorted([alvo for alvo in df_alvos_brutos['target_name'].unique().tolist() if alvo.strip() != ""]) if not df_alvos_brutos.empty else []
                        alvo_sel = st.selectbox(" 2. Selecione o Alvo", ["TODOS OS ALVOS"] + lista_alvos)

                        st.markdown("<br>", unsafe_allow_html=True)
                        
                        objetivo_texto = "Não informado"
                        if not df_p_view.empty and 'objective' in df_p_view.columns:
                            val_bhave = df_p_view['objective'].iloc[0]
                            if pd.notna(val_bhave) and str(val_bhave).strip() not in ["", "None", "nan"]:
                                objetivo_texto = val_bhave
                                
                        if objetivo_texto == "Não informado" and not df_lib.empty and prog_sel in df_lib['name'].values:
                            obj_val = df_lib[df_lib['name'] == prog_sel].iloc[0]['objective_template']
                            if pd.notna(obj_val) and str(obj_val).strip() != "":
                                objetivo_texto = obj_val
                                
                        st.markdown(f"** Meta do Programa:** {objetivo_texto}")
                        st.markdown("###  Desempenho no Período Selecionado")

                        if not df_p_view.empty:
                            if 'phase' not in df_p_view.columns: df_p_view['phase'] = ""
                            
                            mask_lb_prog = df_p_view['phase'].astype(str).str.contains(r'linha de base|\blb\b|baseline|sondagem', case=False, na=False)
                            df_p_int = df_p_view[~mask_lb_prog] 
                            df_p_lb = df_p_view[mask_lb_prog]   
                            
                            media_indep_int = df_p_int['independent_rate'].mean() if not df_p_int.empty else None
                            media_indep_lb = df_p_lb['independent_rate'].mean() if not df_p_lb.empty else None
                            
                            ultima_data = df_p_view['date'].iloc[-1]
                            data_formatada = pd.to_datetime(ultima_data).strftime("%d/%m/%Y")
                            
                            col_m1, col_m2, col_m3 = st.columns(3)
                            
                            col_m1.metric(" Última Aplicação", data_formatada)
                            
                            delta_lb = f"{media_indep_int - media_indep_lb:.1f}% vs LB" if (pd.notnull(media_indep_int) and pd.notnull(media_indep_lb)) else None
                            col_m2.metric(" Indep. Média (Intervenção)", f"{media_indep_int:.1f}%" if pd.notnull(media_indep_int) else "N/A", delta=delta_lb)
                            col_m3.metric(" Média Linha de Base", f"{media_indep_lb:.1f}%" if pd.notnull(media_indep_lb) else "N/A")

                            st.markdown("---")
                            col_geral, col_especifica = st.columns([1, 1])

                            with col_geral:
                                st.subheader(" Gráfico de Colunas (Sessões)")
                                
                                df_p_view['data_str'] = pd.to_datetime(df_p_view['date']).dt.strftime('%d/%m/%Y')
                                df_p_view['phase_clean'] = df_p_view['phase'].astype(str).str.upper().str.strip()
                                df_p_view['phase_clean'] = df_p_view['phase_clean'].replace(["NAN", "NONE", ""], "NÃO INFORMADA")

                                # AGRUPAMENTO DIÁRIO PARA EVITAR DUPLICIDADE (QUADRADOS DUPLOS NO GRÁFICO)
                                df_chart = df_p_view.groupby(['data_str', 'date'], as_index=False).agg({
                                    'independent_rate': 'mean',
                                    'prompt_rate': 'mean',
                                    'phase_clean': 'last'
                                }).sort_values('date')

                                df_chart['Independentes'] = (df_chart['independent_rate'] * 10 / 100).round(0)
                                df_chart['Com Ajuda'] = (df_chart['prompt_rate'] * 10 / 100).round(0)
                                df_chart['Erros'] = (10 - (df_chart['Independentes'] + df_chart['Com Ajuda'])).clip(lower=0)

                                fases_mudanca = []
                                last_valid_phase = None
                                for i, row_fase in df_chart.iterrows():
                                    curr_phase = str(row_fase['phase_clean'])
                                    if curr_phase not in ["NÃO INFORMADA"]:
                                        if last_valid_phase is not None and curr_phase != last_valid_phase:
                                            fases_mudanca.append({'data_str': row_fase['data_str'], 'fase': curr_phase})
                                        last_valid_phase = curr_phase

                                df_p_melt = df_chart.melt(id_vars=['data_str'], value_vars=['Independentes', 'Com Ajuda', 'Erros'], var_name='Status', value_name='Quantidade')
                                fig_p = px.bar(df_p_melt, x='data_str', y='Quantidade', color='Status', 
                                               color_discrete_map={'Independentes': '#2ECC71', 'Com Ajuda': '#F1C40F', 'Erros': '#E74C3C'})
                                
                                fig_p.update_yaxes(range=[0, 11])
                                fig_p.update_xaxes(type='category', tickangle=-45)
                                
                                for mudanca in fases_mudanca:
                                    fig_p.add_vline(x=mudanca['data_str'], line_width=2, line_dash="dash", line_color="black")
                                    fig_p.add_annotation(
                                        x=mudanca['data_str'], y=10.5, 
                                        text=f"  {mudanca['fase']} ",
                                        showarrow=False, textangle=-90, yanchor="bottom",
                                        font=dict(color="black", size=11, weight="bold")
                                    )
                                    
                                st.plotly_chart(fig_p, use_container_width=True)

                            with col_especifica:
                                if alvo_sel == "TODOS OS ALVOS":
                                    st.subheader(" Desempenho por Alvo (%)")
                                    if not df_alvos_brutos.empty:
                                        if 'phase' not in df_alvos_brutos.columns: df_alvos_brutos['phase'] = ""
                                        mask_a_lb_geral = df_alvos_brutos['phase'].astype(str).str.contains(r'linha de base|\blb\b|baseline|sondagem', case=False, na=False)
                                        df_alvos_int = df_alvos_brutos[~mask_a_lb_geral]
                                        
                                        if not df_alvos_int.empty:
                                            df_res_alvos = df_alvos_int.groupby('target_name')['independent_rate'].mean().reset_index().sort_values('independent_rate', ascending=False)
                                            fig_bar_alvos = px.bar(df_res_alvos, x='target_name', y='independent_rate', text_auto='.1f', color='independent_rate', color_continuous_scale='RdYlGn', range_y=[0, 100])
                                            st.plotly_chart(fig_bar_alvos, width='stretch')
                                        else:
                                            st.info(" Todos os alvos encontram-se em Linha de Base (Aguardando intervenção).")
                                else:
                                    st.subheader(f" Detalhe: {alvo_sel}")
                                    df_a_view = df_alvos_brutos[df_alvos_brutos['target_name'] == alvo_sel].sort_values('date').copy()
                                    if not df_a_view.empty:
                                        df_a_view['data_str'] = pd.to_datetime(df_a_view['date']).dt.strftime('%d/%m/%Y')
                                        
                                        # Agrupamento também nos alvos para evitar barras duplas
                                        df_a_chart = df_a_view.groupby(['data_str', 'date'], as_index=False).agg({
                                            'independent_rate': 'mean',
                                            'prompt_rate': 'mean'
                                        }).sort_values('date')

                                        df_a_chart['Independentes'] = (df_a_chart['independent_rate'] * 10 / 100).round(0)
                                        df_a_chart['Com Ajuda'] = (df_a_chart['prompt_rate'] * 10 / 100).round(0)
                                        df_a_chart['Erros'] = (10 - (df_a_chart['Independentes'] + df_a_chart['Com Ajuda'])).clip(lower=0)
                                        df_melt_a = df_a_chart.melt(id_vars=['data_str'], value_vars=['Independentes', 'Com Ajuda', 'Erros'], var_name='Status', value_name='Qtd')
                                        
                                        fig_a = px.bar(df_melt_a, x='data_str', y='Qtd', color='Status', color_discrete_map={'Independentes': '#2ECC71', 'Com Ajuda': '#F1C40F', 'Erros': '#E74C3C'})
                                        fig_a.update_xaxes(type='category', tickangle=-45)
                                        st.plotly_chart(fig_a, width='stretch')
                               
            # ==========================================
            # --- ABA 2: DECISÃO CLÍNICA ---
            # ==========================================
            with tab2:
                st.markdown(f"### Motor de Decisão Clínica: **{paciente_sel}**")
                df_p_filt = pd.DataFrame()
                df_b_decisao = pd.DataFrame()
                if not df_p_raw.empty:
                    min_d_d, max_d_d = df_p_raw['date'].min(), df_p_raw['date'].max()
                    c_dt_d, _ = st.columns([1, 2])
                    periodo_d = c_dt_d.date_input("Período de Análise", value=(min_d_d, max_d_d), format="DD/MM/YYYY", key="decisao_periodo")
                    start_d, end_d = (periodo_d[0], periodo_d[1]) if len(periodo_d) == 2 else (min_d_d, max_d_d)
                    df_p_filt = df_p_raw[(df_p_raw['date'] >= start_d) & (df_p_raw['date'] <= end_d)]
                    if not df_b_raw.empty:
                        df_b_decisao = df_b_raw[(df_b_raw['date'] >= start_d) & (df_b_raw['date'] <= end_d)].copy()
                
                with st.expander("Critério Padrão da Clínica (Global)"):
                    col_cfg1, col_cfg2 = st.columns(2)
                    global_pct = col_cfg1.number_input("Independência (%)", value=90, key="glob_pct")
                    global_dias = col_cfg2.number_input("Dias Consecutivos", value=10, key="glob_dias")

                st.markdown("---")
                if df_p_filt.empty:
                    st.warning("Sem dados no período selecionado.")
                else:
                    progs_ativos = sorted(df_p_filt['programa'].unique().tolist())
                    prog_decisao = st.selectbox("Selecione o Programa para Análise Detalhada", ["RESUMO GERAL"] + progs_ativos)

                    if prog_decisao == "RESUMO GERAL":
                        resumo_decisao = []
                        for p in progs_ativos:
                            df_h = df_p_filt[df_p_filt['programa'] == p].sort_values('date')
                            m_pct, m_dias = global_pct, global_dias
                            if not df_lib.empty and p in df_lib['name'].values:
                                reg = df_lib[df_lib['name'] == p].iloc[0]
                                m_pct, m_dias = int(reg['mastery_threshold_percent']), int(reg['mastery_days'])
                            streak = 0
                            for rate in df_h['independent_rate']:
                                if rate >= m_pct: streak += 1
                                else: streak = 0
                            status = "Atingido" if streak >= m_dias else "Em Evolução" if streak > 0 else "Estagnado"
                            resumo_decisao.append({"Programa": p, "Status": status, "Sessões Consecutivas": streak, "Meta": f"{m_pct}% / {m_dias}d"})
                        st.subheader("Status dos Objetivos")
                        st.dataframe(pd.DataFrame(resumo_decisao).style.map(
                            lambda val: 'background-color: #d4edda; color: #155724;' if val == 'Atingido' else 'background-color: #fff3cd; color: #856404;' if val == 'Em Evolução' else 'background-color: #f8d7da; color: #721c24;' if val == 'Estagnado' else '', subset=['Status']
                        ), width='stretch', hide_index=True)
                    else:
                        st.subheader(f"Análise Detalhada: {prog_decisao}")
                        m_pct, m_dias = global_pct, global_dias
                        if not df_lib.empty and prog_decisao in df_lib['name'].values:
                            reg = df_lib[df_lib['name'] == prog_decisao].iloc[0]
                            m_pct, m_dias = int(reg['mastery_threshold_percent']), int(reg['mastery_days'])
                        
                        c_adj1, c_adj2 = st.columns(2)
                        adj_pct = c_adj1.number_input(f"Ajustar Independência %", value=m_pct, key="adj_pct")
                        adj_dias = c_adj2.number_input(f"Ajustar Dias Consecutivos", value=m_dias, key="adj_dias")

                        df_alvos = load_targets_from_api(paciente_sel, prog_decisao)
                        df_alvos_filt = df_alvos[(df_alvos['date'] >= start_d) & (df_alvos['date'] <= end_d)] if not df_alvos.empty else pd.DataFrame()
                        
                        if not df_alvos_filt.empty:
                            fases_map = df_p_raw[df_p_raw['programa'] == prog_decisao][['date', 'phase']].set_index('date')['phase'].to_dict()
                            status_alvos = []
                            for alvo in sorted(df_alvos_filt['target_name'].unique()):
                                df_a_s = df_alvos_filt[df_alvos_filt['target_name'] == alvo].sort_values('date')
                                stk = 0
                                sessoes_lb = 0
                                sessoes_ajuda = 0
                                for index, row in df_a_s.iterrows():
                                    fase_orig = str(fases_map.get(row['date'], "")).lower().strip()
                                    is_lb = any(termo in fase_orig for termo in ["linha de base", "baseline", "sondagem"]) or bool(re.search(r'\blb\b', fase_orig))
                                    if is_lb: sessoes_lb += 1
                                    else: sessoes_ajuda += 1
                                    if row['independent_rate'] >= adj_pct: stk += 1
                                    else: stk = 0
                                res_a = "Independente" if stk >= adj_dias else "Em Treino"
                                status_alvos.append({
                                    "Alvo": alvo, "Status": res_a, "Sessões em LB": sessoes_lb, 
                                    "Sessões em Ensino": sessoes_ajuda, "Seq. Independência": stk, 
                                    "Última Indep.": f"{df_a_s['independent_rate'].iloc[-1]}%"
                                })
                            st.table(pd.DataFrame(status_alvos))

                st.markdown("---")
                st.subheader("Interferentes no periodo de analise")
                if df_b_decisao.empty:
                    st.info("Sem comportamentos interferentes registrados para o periodo selecionado.")
                else:
                    df_b_decisao['date_pd'] = pd.to_datetime(df_b_decisao['date'])
                    df_b_decisao['data_str'] = df_b_decisao['date_pd'].dt.strftime('%d/%m/%Y')
                    resumo_interferentes = (
                        df_b_decisao.groupby('comportamento', as_index=False)
                        .agg(
                            total_ocorrencias=('count', 'sum'),
                            taxa_media=('rate', 'mean'),
                            taxa_maxima=('rate', 'max'),
                            registros=('date', 'count'),
                            ultima_data=('date_pd', 'max'),
                        )
                        .sort_values(['total_ocorrencias', 'taxa_media'], ascending=False)
                    )

                    total_ocorrencias = resumo_interferentes['total_ocorrencias'].sum()
                    taxa_media_periodo = df_b_decisao['rate'].mean()
                    comportamento_principal = resumo_interferentes['comportamento'].iloc[0]
                    col_int_1, col_int_2, col_int_3 = st.columns(3)
                    col_int_1.metric("Ocorrencias", f"{total_ocorrencias:.0f}")
                    col_int_2.metric("Comportamentos", df_b_decisao['comportamento'].nunique())
                    col_int_3.metric("Taxa media", f"{taxa_media_periodo:.2f}" if pd.notnull(taxa_media_periodo) else "N/A")

                    st.info(
                        "Leitura inicial: o comportamento com maior peso no periodo foi "
                        f"{comportamento_principal}. Considere cruzar os dias de pico com demandas, fases de ensino, "
                        "nivel de ajuda e eventos antecedentes/consequentes antes de decidir proximos passos."
                    )

                    col_graf_int_1, col_graf_int_2 = st.columns(2)
                    with col_graf_int_1:
                        fig_int_count = px.bar(
                            df_b_decisao.sort_values('date_pd'),
                            x='data_str',
                            y='count',
                            color='comportamento',
                            barmode='group',
                            title="Interferentes por sessao",
                            labels={'data_str': 'Data', 'count': 'Contagem', 'comportamento': 'Comportamento'},
                        )
                        fig_int_count.update_xaxes(type='category', tickangle=-45)
                        st.plotly_chart(fig_int_count, use_container_width=True)
                    with col_graf_int_2:
                        fig_int_rate = px.line(
                            df_b_decisao.sort_values('date_pd'),
                            x='date_pd',
                            y='rate',
                            color='comportamento',
                            markers=True,
                            title="Taxa de interferentes",
                            labels={'date_pd': 'Data', 'rate': 'Taxa', 'comportamento': 'Comportamento'},
                        )
                        fig_int_rate.update_xaxes(type='date', tickformat="%d/%m/%y", tickangle=-45)
                        st.plotly_chart(fig_int_rate, use_container_width=True)

                    resumo_interferentes['ultima_data'] = resumo_interferentes['ultima_data'].dt.strftime('%d/%m/%Y')
                    st.dataframe(
                        resumo_interferentes.rename(columns={
                            'comportamento': 'Comportamento',
                            'total_ocorrencias': 'Ocorrencias',
                            'taxa_media': 'Taxa media',
                            'taxa_maxima': 'Taxa maxima',
                            'registros': 'Registros',
                            'ultima_data': 'Ultima data',
                        }),
                        use_container_width=True,
                        hide_index=True,
                    )

            # ==========================================
            # --- ABA 3: INTERFERENTES ---
            # ==========================================
            with tab3:
                st.markdown(f"###  Manejo de Crises: **{paciente_sel}**")
                if not df_b_raw.empty and not df_b_raw['date'].dropna().empty:
                    min_d_b, max_d_b = df_b_raw['date'].min(), df_b_raw['date'].max()
                    c_dt_b, _ = st.columns([1, 2])
                    periodo_b = c_dt_b.date_input(" Período (Interferentes)", value=(min_d_b, max_d_b), format="DD/MM/YYYY", key="periodo_b")
                    st_b, en_b = (periodo_b[0], periodo_b[1]) if len(periodo_b) == 2 else (min_d_b, max_d_b)
                    df_beh = df_b_raw[(df_b_raw['date'] >= st_b) & (df_b_raw['date'] <= en_b)]
                    
                    if df_beh.empty:
                        st.info("Sem comportamentos interferentes para este período.")
                    else:
                        beh_sel = st.selectbox("Selecione o Comportamento", ["VISÃO GERAL (Todos)"] + sorted(df_beh['comportamento'].unique().tolist()))
                        if beh_sel == "VISÃO GERAL (Todos)":
                            c1, c2 = st.columns(2)
                            with c1:
                                fig1 = px.bar(df_beh, x='date', y='count', color='comportamento', barmode='group', title="Frequência (Contagem)")
                                fig1.update_xaxes(type='category', tickangle=-45) 
                                st.plotly_chart(fig1, use_container_width=True)
                            with c2:
                                fig2 = px.line(df_beh, x='date', y='rate', color='comportamento', markers=True, title="Taxa (Ocorrências/Hora)")
                                fig2.update_xaxes(type='date', tickformat="%d/%m/%y", tickangle=-45)
                                st.plotly_chart(fig2, use_container_width=True)
                        else:
                            df_f = df_beh[df_beh['comportamento'] == beh_sel].sort_values('date')
                            fig = px.line(df_f, x='date', y='rate', markers=True, title=f"Evolução: {beh_sel}")
                            fig.update_xaxes(type='date', tickformat="%d/%m/%y", tickangle=-45)
                            st.plotly_chart(fig, use_container_width=True)

            # ==========================================
            # --- ABA 4: RELATÓRIO PEI ---
            # ==========================================
            with tab4:
                import matplotlib.ticker as ticker
                
                st.markdown(f"### Relatorio PEI: **{paciente_sel}**")
                
                START_DATE = datetime(2026, 3, 27)
                ref_date = START_DATE.date()
                hoje = datetime.now().date()

                periodo_pei = st.date_input(
                    "Periodo do PEI",
                    value=(max(ref_date, hoje - timedelta(days=90)), hoje if hoje >= ref_date else ref_date),
                    min_value=ref_date,
                    max_value=hoje,
                    format="DD/MM/YYYY",
                    key="pei_periodo",
                )
                if isinstance(periodo_pei, tuple) and len(periodo_pei) == 2:
                    ciclo_inicio_pei, ciclo_fim_pei = periodo_pei
                else:
                    ciclo_inicio_pei, ciclo_fim_pei = ref_date, hoje
                if ciclo_inicio_pei > ciclo_fim_pei:
                    ciclo_inicio_pei, ciclo_fim_pei = ciclo_fim_pei, ciclo_inicio_pei
                ciclo_fim_exclusivo_pei = ciclo_fim_pei + timedelta(days=1)
                st.caption(
                    "Periodo selecionado: "
                    f"{ciclo_inicio_pei.strftime('%d/%m/%Y')} a {ciclo_fim_pei.strftime('%d/%m/%Y')}"
                )

                programas_pei = sorted([
                    str(programa).strip()
                    for programa in df_p_raw.get("programa", pd.Series(dtype=str)).dropna().unique()
                    if str(programa).strip()
                ])
                objetivo_grafico_pei = st.selectbox(
                    "Objetivo para o grafico do PEI",
                    ["Media geral dos objetivos"] + programas_pei,
                    key="pei_objetivo_grafico",
                )
                usar_ia_texto_pei = st.checkbox(
                    "Gerar texto do periodo com IA",
                    value=True,
                    key="pei_usar_ia_texto",
                )

                def verificar_alvos_clean(objetivo_texto):
                    if not objetivo_texto or str(objetivo_texto).strip() == "" or str(objetivo_texto).lower() in ["nan", "none"]:
                        return "Descrição não informada."
                    return str(objetivo_texto)

                def calcular_evolucao_trimestral(desempenho_anterior, desempenho_atual, criterio=90):
                    if pd.isna(desempenho_atual): return "-"
                    if desempenho_atual >= criterio: return "Atingiu (ATG)"
                    if pd.isna(desempenho_anterior): return "Em Avaliação"
                    if desempenho_atual > desempenho_anterior + 5: return "Avançou (AVA)"
                    if desempenho_atual < desempenho_anterior - 5: return "Agravou (AGR)"
                    return "Estabilizou (EST)"

                def limpar_texto_pei(texto):
                    texto = str(texto)
                    substituicoes = {
                        "": "",
                        "": "",
                        "": "",
                        "": "",
                        "": "",
                        "": "",
                        "": "",
                        "": "",
                        "": "",
                        "•": "-",
                    }
                    for antigo, novo in substituicoes.items():
                        texto = texto.replace(antigo, novo)
                    return re.sub(r"[\U0001F300-\U0001FAFF\u2600-\u27BF]\ufe0f?", "", texto).strip()

                def resposta_ia_invalida(texto):
                    texto_limpo = str(texto or "").strip()
                    texto_cf = texto_limpo.casefold()
                    marcadores_bloqueados = [
                        "```",
                        "{",
                        "}",
                        "[",
                        "]",
                        "json",
                        "python",
                        "summary",
                        "behavior",
                        "patient_context",
                        "fontes_recuperadas",
                        "dados_do_paciente",
                        "response",
                    ]
                    return not texto_limpo or any(marcador in texto_cf for marcador in marcadores_bloqueados)

                def filtrar_periodo_pei(df, inicio, fim_exclusivo):
                    if df is None or df.empty or "date" not in df.columns:
                        return pd.DataFrame() if df is None else df.copy()
                    df_periodo = df.copy()
                    df_periodo["date_pd"] = pd.to_datetime(df_periodo["date"], errors="coerce")
                    df_periodo = df_periodo.dropna(subset=["date_pd"])
                    return df_periodo[
                        (df_periodo["date_pd"] >= pd.Timestamp(inicio))
                        & (df_periodo["date_pd"] < pd.Timestamp(fim_exclusivo))
                    ].copy()

                def carregar_alvos_programas_pei(programas):
                    todos_alvos = []
                    for prog in programas:
                        alvos = load_targets_from_api(paciente_sel, prog)
                        if not alvos.empty:
                            alvos = alvos.copy()
                            alvos["programa"] = prog
                            todos_alvos.append(alvos)
                    return pd.concat(todos_alvos, ignore_index=True) if todos_alvos else pd.DataFrame()

                def limiar_programa_pei(programa, df_lib, padrao=90):
                    if df_lib.empty or "name" not in df_lib.columns or not programa:
                        return padrao
                    registros = df_lib[df_lib["name"] == programa]
                    if registros.empty:
                        return padrao
                    valor = registros.iloc[0].get("mastery_threshold_percent", padrao)
                    return float(valor) if pd.notna(valor) else padrao

                def status_por_independencia(inicial, atual, limiar=90):
                    if pd.isna(atual):
                        return "Sem registro"
                    if atual >= limiar:
                        return "Atingiu"
                    if pd.isna(inicial):
                        return "Em avaliação"
                    if atual > inicial + 5:
                        return "Avançou"
                    if atual < inicial - 5:
                        return "Agravou"
                    return "Estabilizou"

                def preparar_alvos_programa(df_alvos, programa):
                    if (
                        df_alvos is None
                        or df_alvos.empty
                        or "programa" not in df_alvos.columns
                        or "target_name" not in df_alvos.columns
                    ):
                        return pd.DataFrame()
                    df_prog_alvos = df_alvos[df_alvos["programa"] == programa].copy()
                    if df_prog_alvos.empty:
                        return pd.DataFrame()
                    if "date_pd" not in df_prog_alvos.columns and "date" in df_prog_alvos.columns:
                        df_prog_alvos["date_pd"] = pd.to_datetime(df_prog_alvos["date"], errors="coerce")
                    if "independent_rate" in df_prog_alvos.columns:
                        df_prog_alvos["independent_rate"] = pd.to_numeric(df_prog_alvos["independent_rate"], errors="coerce")
                    if "phase" not in df_prog_alvos.columns:
                        df_prog_alvos["phase"] = ""
                    return df_prog_alvos

                def resumo_alvos_programa(df_alvos, programa, df_lib, df_alvos_periodo=None):
                    colunas = ["Alvo", "LinhaBase", "Periodo", "Resumo", "Status"]
                    df_prog_alvos = preparar_alvos_programa(df_alvos, programa)
                    df_prog_periodo = preparar_alvos_programa(df_alvos_periodo if df_alvos_periodo is not None else df_alvos, programa)
                    if df_prog_periodo.empty:
                        return pd.DataFrame(columns=colunas)

                    limiar = limiar_programa_pei(programa, df_lib)
                    rows = []
                    for alvo, grupo_periodo in df_prog_periodo.groupby("target_name"):
                        grupo_todos = df_prog_alvos[df_prog_alvos["target_name"] == alvo].copy()
                        if grupo_todos.empty:
                            grupo_todos = grupo_periodo.copy()
                        grupo_todos = grupo_todos.sort_values("date_pd") if "date_pd" in grupo_todos.columns else grupo_todos.copy()
                        grupo_periodo = grupo_periodo.sort_values("date_pd") if "date_pd" in grupo_periodo.columns else grupo_periodo.copy()

                        serie_todos = grupo_todos["independent_rate"].dropna() if "independent_rate" in grupo_todos.columns else pd.Series(dtype=float)
                        serie_periodo = grupo_periodo["independent_rate"].dropna() if "independent_rate" in grupo_periodo.columns else pd.Series(dtype=float)
                        mask_lb = grupo_todos["phase"].astype(str).str.contains(r"linha de base|\blb\b|baseline|sondagem", case=False, na=False)
                        serie_lb = grupo_todos.loc[mask_lb, "independent_rate"].dropna() if "independent_rate" in grupo_todos.columns else pd.Series(dtype=float)
                        inicial = serie_lb.mean() if not serie_lb.empty else (serie_todos.iloc[0] if not serie_todos.empty else pd.NA)
                        atual = serie_periodo.iloc[-1] if not serie_periodo.empty else pd.NA
                        media_ind = serie_periodo.mean() if not serie_periodo.empty else pd.NA
                        linha_base_txt = f"{inicial:.1f}% de independência na linha de base" if pd.notna(inicial) else "linha de base sem registro percentual"
                        periodo_txt = f"{media_ind:.1f}% de independência no período observado" if pd.notna(media_ind) else "período observado sem registro percentual"
                        rows.append({
                            "Alvo": limpar_texto_pei(alvo),
                            "LinhaBase": linha_base_txt,
                            "Periodo": periodo_txt,
                            "Resumo": f"{linha_base_txt}; {periodo_txt}",
                            "Status": status_por_independencia(inicial, atual, limiar),
                        })
                    return pd.DataFrame(rows, columns=colunas).sort_values(["Status", "Alvo"]).reset_index(drop=True)

                def resumo_objetivos_por_area(areas):
                    rows = []
                    for area_num, itens in areas.items():
                        for item_index, item in enumerate(itens):
                            rows.append({
                                "Objetivo": f"Objetivo {area_num}.{item_index + 1}",
                                "Descrição": limpar_texto_pei(item.get("texto", "")),
                            })
                        if not itens:
                            rows.append({
                                "Objetivo": "-",
                                "Descrição": "Sem objetivos cadastrados para esta área.",
                            })
                    return pd.DataFrame(rows, columns=["Objetivo", "Descrição"])

                def formatar_lista_alvos_corrida(alvos):
                    if alvos.empty:
                        return "Sem alvos trabalhados no período"
                    partes = []
                    for _, alvo_row in alvos.iterrows():
                        alvo = limpar_texto_pei(alvo_row.get("Alvo", ""))
                        resumo = limpar_texto_pei(alvo_row.get("Resumo", ""))
                        partes.append(f"{alvo} ({resumo})")
                    return "; ".join(partes)

                def status_geral_alvos(alvos):
                    if alvos.empty:
                        return "Sem registro"
                    prioridades = {"Agravou": 0, "Estabilizou": 1, "Em avaliação": 2, "Avançou": 3, "Atingiu": 4}
                    return sorted(
                        alvos["Status"].tolist(),
                        key=lambda status: prioridades.get(str(status), 2),
                        reverse=True,
                    )[0]

                def periodo_aplicacao_programa(df_hist, programa, periodo_inicio, periodo_fim, dias_inatividade=30):
                    if df_hist is None or df_hist.empty or not programa or "programa" not in df_hist.columns or "date" not in df_hist.columns:
                        return f"{periodo_inicio.strftime('%d/%m/%Y')} a {periodo_fim.strftime('%d/%m/%Y')}"
                    hist = df_hist[df_hist["programa"] == programa].copy()
                    if hist.empty:
                        return f"{periodo_inicio.strftime('%d/%m/%Y')} a {periodo_fim.strftime('%d/%m/%Y')}"
                    hist["date_pd"] = pd.to_datetime(hist["date"], errors="coerce")
                    hist = hist.dropna(subset=["date_pd"])
                    if hist.empty:
                        return f"{periodo_inicio.strftime('%d/%m/%Y')} a {periodo_fim.strftime('%d/%m/%Y')}"

                    inicio_ts = pd.Timestamp(periodo_inicio)
                    fim_ts = pd.Timestamp(periodo_fim)
                    hist_periodo = hist[(hist["date_pd"] >= inicio_ts) & (hist["date_pd"] <= fim_ts)]
                    if hist_periodo.empty:
                        return f"{periodo_inicio.strftime('%d/%m/%Y')} a {periodo_fim.strftime('%d/%m/%Y')}"

                    primeira = hist_periodo["date_pd"].min()
                    ultima = hist_periodo["date_pd"].max()
                    inicio_real = max(inicio_ts, primeira)
                    fim_real = ultima if ultima < (fim_ts - pd.Timedelta(days=dias_inatividade)) else fim_ts
                    return f"{inicio_real.strftime('%d/%m/%Y')} a {fim_real.strftime('%d/%m/%Y')}"

                def texto_desempenho_inicial(valor):
                    return f"{valor:.1f}% de independência na linha de base" if pd.notna(valor) else "Linha de base sem registro percentual"

                def resumo_alvos_por_objetivo(areas, df_hist, df_alvos, df_alvos_periodo, df_lib, periodo_inicio, periodo_fim):
                    rows = []
                    for area_num, itens in areas.items():
                        for item_index, item in enumerate(itens):
                            codigo = f"Objetivo {area_num}.{item_index + 1}"
                            programa = item.get("programa", "")
                            limiar = limiar_programa_pei(programa, df_lib)
                            linha_de_base, _ = desempenho_trimestral_para_programa(programa, df_hist, limiar)
                            desempenho_inicial = texto_desempenho_inicial(linha_de_base)
                            periodo = periodo_aplicacao_programa(df_hist, programa, periodo_inicio, periodo_fim)
                            alvos = resumo_alvos_programa(df_alvos, programa, df_lib, df_alvos_periodo)
                            rows.append({
                                "Objetivo": codigo,
                                "Período avaliado": periodo,
                                "Desempenho inicial": desempenho_inicial,
                                "Alvos trabalhados": formatar_lista_alvos_corrida(alvos),
                                "Status": status_geral_alvos(alvos),
                            })
                    return pd.DataFrame(rows, columns=["Objetivo", "Período avaliado", "Desempenho inicial", "Alvos trabalhados", "Status"])

                def aplicar_fonte_pei(doc):
                    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
                        try:
                            style = doc.styles[style_name]
                            style.font.name = "Times New Roman"
                            style.font.size = Pt(12)
                            style.font.color.rgb = RGBColor(0, 0, 0)
                            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        except Exception:
                            pass

                    def formatar_paragrafo(paragraph):
                        for run in paragraph.runs:
                            if run._element.xpath(".//*[local-name()='drawing']") or run._element.xpath(".//*[local-name()='pict']"):
                                continue
                            run.text = limpar_texto_pei(run.text)
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            try:
                                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                            except Exception:
                                pass

                    for paragraph in doc.paragraphs:
                        formatar_paragrafo(paragraph)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    formatar_paragrafo(paragraph)

                def texto_objetivo_programa(row):
                    objetivo = verificar_alvos_clean(row.get("objective", ""))
                    programa = str(row.get("programa", "")).strip()
                    if objetivo and objetivo != "Descrição não informada.":
                        return objetivo
                    return programa or "Objetivo não informado."

                def classificar_area_pei(programa, objetivo):
                    texto = f"{programa} {objetivo}".lower()
                    if any(chave in texto for chave in ["avd", "vida diária", "vida diaria", "higiene", "banheiro", "vestir", "aliment", "escovar", "lavar"]):
                        return 5
                    if any(chave in texto for chave in ["social", "turno", "grupo", "espera", "compartilhar", "interferente", "comportamento", "crise", "fuga"]):
                        return 4
                    if any(chave in texto for chave in ["brinc", "jogo", "brinquedo", "faz de conta", "lúdico", "ludico"]):
                        return 3
                    if any(chave in texto for chave in ["recept", "ouvinte", "instru", "discrimina", "identificar", "apontar", "parear"]):
                        return 2
                    return 1

                def distribuir_objetivos_por_area(df_prog, df_beh):
                    areas = {1: [], 2: [], 3: [], 4: [], 5: []}
                    df_prog_unicos = df_prog.drop_duplicates(subset=["programa"]) if not df_prog.empty else pd.DataFrame()
                    for _, row in df_prog_unicos.iterrows():
                        programa = str(row.get("programa", "")).strip()
                        objetivo = texto_objetivo_programa(row)
                        area = classificar_area_pei(programa, objetivo)
                        areas[area].append({"programa": programa, "texto": objetivo})

                    if not df_beh.empty and "comportamento" in df_beh.columns:
                        comportamentos = sorted([str(valor) for valor in df_beh["comportamento"].dropna().unique() if str(valor).strip()])
                        if comportamentos:
                            texto = (
                                "Reduzir a taxa dos comportamentos interferentes monitorados "
                                f"({', '.join(comportamentos[:8])}) com base nos registros clínicos do período."
                            )
                            areas[4].insert(0, {"programa": "Comportamentos interferentes", "texto": texto})
                    return areas

                def set_cell_text(cell, text):
                    cell.text = limpar_texto_pei(text)

                def aplicar_estilo_tabela_seguro(table, nomes=("Table Grid", "Tabela Grade", "Grade da Tabela")):
                    for nome in nomes:
                        try:
                            table.style = nome
                            return
                        except KeyError:
                            continue
                        except Exception:
                            return

                def aplicar_bordas_grade_tabela(table, color="000000", size="8"):
                    for row in table.rows:
                        for cell in row.cells:
                            tc_pr = cell._tc.get_or_add_tcPr()
                            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
                            if tc_borders is None:
                                tc_borders = OxmlElement("w:tcBorders")
                                tc_pr.append(tc_borders)
                            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                                tag = f"w:{edge}"
                                element = tc_borders.find(qn(tag))
                                if element is None:
                                    element = OxmlElement(tag)
                                    tc_borders.append(element)
                                element.set(qn("w:val"), "single")
                                element.set(qn("w:sz"), size)
                                element.set(qn("w:space"), "0")
                                element.set(qn("w:color"), color)

                def limpar_linhas_tabela(table):
                    while len(table.rows) > 1:
                        table._tbl.remove(table.rows[-1]._tr)
                    for row in table.rows:
                        for cell in row.cells:
                            cell.text = ""

                def garantir_colunas_tabela(table, total_colunas):
                    while len(table.rows[0].cells) < total_colunas:
                        table.add_column(Inches(1.1))

                def deixar_celula_negrito(cell):
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                def aplicar_cor_fundo_celula(cell, color):
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = tc_pr.find(qn("w:shd"))
                    if shd is None:
                        shd = OxmlElement("w:shd")
                        tc_pr.append(shd)
                    shd.set(qn("w:fill"), color)

                def cor_status_pei(status):
                    status_cf = str(status or "").casefold()
                    if "atingiu" in status_cf or "atingido" in status_cf:
                        return "C6EFCE"
                    if "avançou" in status_cf or "avancou" in status_cf:
                        return "BDD7EE"
                    if "estabilizou" in status_cf or "estagnou" in status_cf:
                        return "FFF2CC"
                    if "agravou" in status_cf:
                        return "F4CCCC"
                    return "E7E6E6"

                def cor_status_css_pei(status):
                    return f"background-color: #{cor_status_pei(status)}"

                def estilizar_status_preview_pei(df):
                    styler = df.style
                    if hasattr(styler, "map"):
                        return styler.map(cor_status_css_pei, subset=["Status"])
                    return styler.apply(
                        lambda coluna: [cor_status_css_pei(valor) for valor in coluna],
                        subset=["Status"],
                    )

                def preencher_tabela_objetivos(table, area_num, itens):
                    garantir_colunas_tabela(table, 2)
                    limpar_linhas_tabela(table)

                    headers = ["Objetivo", "Descrição"]
                    for idx, header in enumerate(headers):
                        set_cell_text(table.rows[0].cells[idx], header)
                        deixar_celula_negrito(table.rows[0].cells[idx])

                    for item_index, item in enumerate(itens):
                        cells = table.add_row().cells
                        set_cell_text(cells[0], f"Objetivo {area_num}.{item_index + 1}")
                        set_cell_text(cells[1], item.get("texto", ""))
                        deixar_celula_negrito(cells[0])

                    if not itens:
                        cells = table.add_row().cells
                        set_cell_text(cells[0], "-")
                        set_cell_text(cells[1], "Sem objetivos cadastrados para esta área.")

                    aplicar_estilo_tabela_seguro(table)
                    aplicar_bordas_grade_tabela(table)

                def desempenho_trimestral_para_programa(programa, df_hist, limiar=90):
                    if df_hist.empty or not programa:
                        return None, []
                    hist = df_hist[df_hist["programa"] == programa].copy()
                    if hist.empty:
                        return None, []
                    hist["date_pd"] = pd.to_datetime(hist["date"], errors="coerce")
                    hist = hist.dropna(subset=["date_pd"]).sort_values("date_pd")
                    if hist.empty:
                        return None, []
                    if "phase" not in hist.columns:
                        hist["phase"] = ""
                    mask_lb = hist["phase"].astype(str).str.contains(r"linha de base|\blb\b|baseline|sondagem", case=False, na=False)
                    linha_de_base = hist[mask_lb]["independent_rate"].mean() if not hist[mask_lb].empty else hist["independent_rate"].iloc[0]

                    trimestres = [
                        (START_DATE, START_DATE + timedelta(days=90), "1. DATA"),
                        (START_DATE + timedelta(days=90), START_DATE + timedelta(days=180), "2. DATA"),
                        (START_DATE + timedelta(days=180), START_DATE + timedelta(days=270), "3. DATA"),
                        (START_DATE + timedelta(days=270), START_DATE + timedelta(days=360), "4. DATA"),
                    ]
                    desempenho_anterior = linha_de_base
                    linhas = []
                    for ini, fim, rotulo in trimestres:
                        df_tri = hist[(hist["date_pd"] >= pd.Timestamp(ini)) & (hist["date_pd"] < pd.Timestamp(fim))]
                        if df_tri.empty:
                            linhas.append((rotulo, "-", "Sem avaliações"))
                            continue
                        media = df_tri["independent_rate"].mean()
                        codigo = calcular_evolucao_trimestral(desempenho_anterior, media, limiar)
                        linhas.append((rotulo.replace("DATA", fim.strftime("%d/%m/%Y")), f"{media:.1f}%", codigo))
                        desempenho_anterior = media
                    return linha_de_base, linhas

                def preencher_tabela_desempenho(table, area_num, itens, df_hist, df_alvos, df_alvos_periodo, df_lib, periodo_inicio, periodo_fim):
                    garantir_colunas_tabela(table, 4)
                    limpar_linhas_tabela(table)

                    headers = ["Objetivo", "Período avaliado", "Desempenho inicial", "Status"]
                    for idx, header in enumerate(headers):
                        set_cell_text(table.rows[0].cells[idx], header)
                        deixar_celula_negrito(table.rows[0].cells[idx])

                    if not itens:
                        periodo = f"{periodo_inicio.strftime('%d/%m/%Y')} a {periodo_fim.strftime('%d/%m/%Y')}"
                        row = table.add_row()
                        set_cell_text(row.cells[0], "-")
                        set_cell_text(row.cells[1], periodo)
                        set_cell_text(row.cells[2], "-")
                        set_cell_text(row.cells[3], "Sem registro")
                        aplicar_cor_fundo_celula(row.cells[3], cor_status_pei("Sem registro"))
                        row_alvos = table.add_row()
                        set_cell_text(row_alvos.cells[0], "Alvos trabalhados")
                        merged = row_alvos.cells[1].merge(row_alvos.cells[3])
                        set_cell_text(merged, "Sem alvos trabalhados no período")
                    else:
                        for item_index, item in enumerate(itens):
                            programa = item.get("programa", "")
                            limiar = limiar_programa_pei(programa, df_lib)
                            linha_de_base, _ = desempenho_trimestral_para_programa(programa, df_hist, limiar)
                            desempenho_inicial = texto_desempenho_inicial(linha_de_base)
                            periodo = periodo_aplicacao_programa(df_hist, programa, periodo_inicio, periodo_fim)
                            codigo = f"Objetivo {area_num}.{item_index + 1}"
                            alvos = resumo_alvos_programa(df_alvos, programa, df_lib, df_alvos_periodo)
                            status_geral = status_geral_alvos(alvos)

                            row = table.add_row()
                            set_cell_text(row.cells[0], codigo)
                            set_cell_text(row.cells[1], periodo)
                            set_cell_text(row.cells[2], desempenho_inicial)
                            set_cell_text(row.cells[3], status_geral)
                            deixar_celula_negrito(row.cells[0])
                            aplicar_cor_fundo_celula(row.cells[3], cor_status_pei(status_geral))

                            row_alvos = table.add_row()
                            set_cell_text(row_alvos.cells[0], "Alvos trabalhados")
                            deixar_celula_negrito(row_alvos.cells[0])
                            merged = row_alvos.cells[1].merge(row_alvos.cells[3])
                            set_cell_text(merged, formatar_lista_alvos_corrida(alvos))
                            aplicar_cor_fundo_celula(merged, cor_status_pei(status_geral))

                    aplicar_estilo_tabela_seguro(table)
                    aplicar_bordas_grade_tabela(table)

                def coletar_evolucoes_periodo(df_hist_periodo, df_alvos_periodo, df_beh_periodo, limite=30):
                    itens = []
                    fontes = [
                        (df_hist_periodo, "programa", "programa"),
                        (df_alvos_periodo, "alvo", "programa"),
                        (df_beh_periodo, "interferente", "comportamento"),
                    ]
                    for df_fonte, tipo, nome_col in fontes:
                        if df_fonte.empty or "evolution" not in df_fonte.columns:
                            continue
                        for _, row in df_fonte.iterrows():
                            texto = limpar_texto_pei(row.get("evolution", ""))
                            if not texto or texto in {"0", "0.0", "nan", "None"}:
                                continue
                            nome = limpar_texto_pei(row.get(nome_col, ""))
                            alvo = limpar_texto_pei(row.get("target_name", "")) if tipo == "alvo" else ""
                            data = row.get("date", "")
                            label = f"{tipo}: {nome}"
                            if alvo:
                                label += f" / {alvo}"
                            itens.append(f"- {label}; data={data}; evolucao={texto}")
                    return "\n".join(itens[:limite])

                def gerar_resumo_trimestral_local(df_hist_periodo, df_alvos_periodo, df_beh_periodo, objetivo_grafico, inicio, fim):
                    periodo = f"{inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}"
                    partes = [f"Analise do periodo de {periodo}."]

                    df_obj = df_hist_periodo.copy()
                    if objetivo_grafico != "Media geral dos objetivos" and not df_obj.empty:
                        df_obj = df_obj[df_obj["programa"] == objetivo_grafico]

                    if not df_obj.empty:
                        media_ind = df_obj["independent_rate"].mean()
                        programas = df_obj["programa"].nunique() if "programa" in df_obj.columns else 0
                        alvo_txt = objetivo_grafico if objetivo_grafico != "Media geral dos objetivos" else f"{programas} objetivos monitorados"
                        trecho = f"No periodo, {alvo_txt} apresentou media de independencia de {media_ind:.1f}%"
                        partes.append(trecho + ".")
                    else:
                        partes.append("Nao houve registros de objetivos no ciclo selecionado.")

                    if not df_alvos_periodo.empty and "target_name" in df_alvos_periodo.columns:
                        total_alvos = df_alvos_periodo["target_name"].nunique()
                        media_alvos = df_alvos_periodo["independent_rate"].mean() if "independent_rate" in df_alvos_periodo.columns else None
                        if pd.notna(media_alvos):
                            partes.append(f"Foram registrados {total_alvos} alvos no periodo, com media de independencia de {media_alvos:.1f}%.")
                        else:
                            partes.append(f"Foram registrados {total_alvos} alvos no periodo.")

                    if not df_beh_periodo.empty and "comportamento" in df_beh_periodo.columns:
                        total = df_beh_periodo["count"].sum() if "count" in df_beh_periodo.columns else 0
                        taxa = df_beh_periodo["rate"].mean() if "rate" in df_beh_periodo.columns else None
                        principais = ", ".join(
                            df_beh_periodo.groupby("comportamento")["count"].sum().sort_values(ascending=False).head(3).index.astype(str)
                        ) if "count" in df_beh_periodo.columns else ", ".join(df_beh_periodo["comportamento"].dropna().astype(str).unique()[:3])
                        trecho = f"Os comportamentos interferentes somaram {total:.0f} ocorrencias"
                        if pd.notna(taxa):
                            trecho += f", com taxa media de {taxa:.2f}"
                        if principais:
                            trecho += f". Principais registros: {principais}"
                        partes.append(trecho + ".")
                    else:
                        partes.append("Nao houve comportamentos interferentes registrados no ciclo selecionado.")

                    evolucoes = coletar_evolucoes_periodo(df_hist_periodo, df_alvos_periodo, df_beh_periodo, limite=12)
                    if evolucoes:
                        partes.append("Evolucoes registradas no periodo:\n" + evolucoes)

                    return "\n\n".join(partes)

                def gerar_texto_trimestral_pei(nome_paciente, df_hist_periodo, df_alvos_periodo, df_beh_periodo, objetivo_grafico, inicio, fim):
                    resumo_local = gerar_resumo_trimestral_local(
                        df_hist_periodo, df_alvos_periodo, df_beh_periodo, objetivo_grafico, inicio, fim
                    )
                    if not usar_ia_texto_pei:
                        return resumo_local

                    pergunta = (
                        "Responda exclusivamente em portugues brasileiro, em paragrafo corrido. "
                        "Nao use codigo, JSON, markdown, listas tecnicas, nomes de variaveis ou termos em ingles. "
                        "Responda apenas com o texto final do PEI, sem repetir estas instrucoes. "
                        "Escreva um texto tecnico, objetivo e cauteloso para inserir no PEI. "
                        f"O periodo selecionado e de {inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}. "
                        "Leia e sintetize as evolucoes/anotacoes registradas pelos terapeutas no periodo, "
                        "sem citar nomes de terapeutas. "
                        "Explique como foram os objetivos no periodo selecionado, destacando "
                        "progresso, estabilidade, agravamento, alvos relevantes e comportamentos interferentes. "
                        "Use apenas os dados e evolucoes do periodo e nao invente informacoes. "
                        f"Objetivo em destaque para o grafico: {objetivo_grafico}."
                    )
                    def parece_prompt_ia(texto):
                        texto_cf = str(texto or "").casefold()
                        marcadores_prompt = [
                            "responda apenas com o texto final do pei",
                            "escreva um texto tecnico",
                            "explique como foram os objetivos",
                            "objetivo em destaque para o grafico",
                            "voce e um assistente especialista",
                            "diretrizes de comportamento",
                            "fontes_recuperadas_nao_confiaveis",
                            "dados_do_paciente_nao_confiaveis",
                        ]
                        return any(marcador in texto_cf for marcador in marcadores_prompt)

                    try:
                        resposta = ask_clinical_agent(nome_paciente, pergunta, start_date=inicio, end_date=fim)
                        modo = str(resposta.get("modo", "")).casefold()
                        if modo in {"busca_local", "busca_local_fallback", "limite_ia", "sem_resultados"}:
                            return resumo_local
                        texto = limpar_texto_pei(resposta.get("resposta", ""))
                        if texto and not parece_prompt_ia(texto) and not resposta_ia_invalida(texto):
                            return texto
                    except Exception as exc:
                        st.warning(f"Nao foi possivel gerar o texto com IA. Usei um resumo automatico local. Detalhe: {exc}")
                    return resumo_local

                def gerar_resumo_comportamentos_problema_local(df_beh_periodo, inicio, fim):
                    periodo = f"{inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}"
                    if df_beh_periodo.empty or "comportamento" not in df_beh_periodo.columns:
                        return f"No periodo de {periodo}, nao houve registros de comportamentos problema."

                    partes = [f"Resumo dos comportamentos problema no periodo de {periodo}."]
                    if "count" in df_beh_periodo.columns:
                        total_por_comp = (
                            df_beh_periodo.groupby("comportamento")["count"]
                            .sum()
                            .sort_values(ascending=False)
                        )
                        if not total_por_comp.empty:
                            principais = "; ".join([f"{comp}: {valor:.0f} ocorrencias" for comp, valor in total_por_comp.head(5).items()])
                            partes.append(f"Maiores registros quantitativos: {principais}.")
                    if "rate" in df_beh_periodo.columns:
                        taxa_por_comp = (
                            df_beh_periodo.groupby("comportamento")["rate"]
                            .mean()
                            .sort_values(ascending=False)
                        )
                        if not taxa_por_comp.empty:
                            taxas = "; ".join([f"{comp}: taxa media {valor:.2f}" for comp, valor in taxa_por_comp.head(5).items()])
                            partes.append(f"Taxas medias observadas: {taxas}.")

                    evolucoes = coletar_evolucoes_periodo(pd.DataFrame(), pd.DataFrame(), df_beh_periodo, limite=10)
                    if evolucoes:
                        partes.append("Evolucoes relevantes registradas no periodo:\n" + evolucoes)
                    else:
                        partes.append("Nao foram encontradas evolucoes textuais especificas para os comportamentos no periodo.")
                    return "\n\n".join(partes)

                def gerar_resumo_comportamentos_problema_pei(nome_paciente, df_beh_periodo, inicio, fim):
                    resumo_local = gerar_resumo_comportamentos_problema_local(df_beh_periodo, inicio, fim)
                    if not usar_ia_texto_pei:
                        return resumo_local

                    evolucoes = coletar_evolucoes_periodo(pd.DataFrame(), pd.DataFrame(), df_beh_periodo, limite=20)
                    pergunta = (
                        "Responda exclusivamente em portugues brasileiro, em paragrafo corrido. "
                        "Nao use codigo, JSON, markdown, listas tecnicas, nomes de variaveis ou termos em ingles. "
                        "Responda apenas com o texto final para o PEI. "
                        "Faca um resumo tecnico, objetivo e cauteloso dos comportamentos problema/interferentes "
                        "do periodo selecionado. Use os registros de evolucao para filtrar o que de fato ocorreu, "
                        "sem inventar informacoes e sem citar nomes de terapeutas. "
                        "Destaque padroes observaveis, comportamentos mais frequentes ou relevantes, possiveis "
                        "contextos mencionados nas evolucoes e impacto sobre os objetivos quando houver dados. "
                        f"Periodo: {inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}. "
                        f"Resumo automatico dos dados: {resumo_local[:1800]} "
                        f"Evolucoes do periodo: {evolucoes[:1500]}"
                    )
                    try:
                        resposta = ask_clinical_agent(nome_paciente, pergunta, start_date=inicio, end_date=fim)
                        modo = str(resposta.get("modo", "")).casefold()
                        if modo in {"busca_local", "busca_local_fallback", "limite_ia", "sem_resultados"}:
                            return resumo_local
                        texto = limpar_texto_pei(resposta.get("resposta", ""))
                        if texto and not resposta_ia_invalida(texto):
                            return texto
                    except Exception as exc:
                        st.warning(f"Nao foi possivel gerar o resumo dos comportamentos com IA. Usei um resumo local. Detalhe: {exc}")
                    return resumo_local

                def adicionar_texto_trimestral(doc, texto, inicio, fim):
                    doc.add_page_break()
                    doc.add_heading("Analise dos objetivos no periodo", level=1)
                    p_periodo = doc.add_paragraph()
                    p_periodo.add_run("Periodo: ").bold = True
                    p_periodo.add_run(f"{inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}")

                    texto_limpo = limpar_texto_pei(texto)
                    if not texto_limpo:
                        texto_limpo = "Espaco reservado para descricao clinica do desempenho dos objetivos no periodo."
                    for bloco in re.split(r"\n\s*\n", texto_limpo):
                        if bloco.strip():
                            doc.add_paragraph(bloco.strip())

                def adicionar_resumo_clinico_periodo(doc, texto, inicio, fim):
                    doc.add_page_break()
                    doc.add_heading("Resumo clínico do período observado", level=1)
                    p_periodo = doc.add_paragraph()
                    p_periodo.add_run("Periodo observado: ").bold = True
                    p_periodo.add_run(f"{inicio.strftime('%d/%m/%Y')} a {fim.strftime('%d/%m/%Y')}")

                    texto_limpo = limpar_texto_pei(texto)
                    if not texto_limpo:
                        texto_limpo = "Sem registros suficientes para sintetizar o período observado."
                    for bloco in re.split(r"\n\s*\n", texto_limpo):
                        if bloco.strip():
                            doc.add_paragraph(bloco.strip())

                def adicionar_grafico_objetivo_periodo(doc, df_hist, objetivo_grafico, inicio, fim_exclusivo):
                    doc.add_heading("Graficos dos objetivos", level=2)
                    df_periodo = filtrar_periodo_pei(df_hist, inicio, fim_exclusivo)

                    if df_periodo.empty:
                        doc.add_paragraph("Sem registros de objetivos no periodo selecionado.")
                        return

                    df_obj = df_periodo.copy()
                    titulo = "Media geral dos objetivos"
                    if objetivo_grafico != "Media geral dos objetivos":
                        df_selecionado = df_periodo[df_periodo["programa"] == objetivo_grafico].copy()
                        if not df_selecionado.empty:
                            df_obj = df_selecionado
                            titulo = objetivo_grafico
                        else:
                            doc.add_paragraph(
                                "O objetivo selecionado nao teve registros no periodo; abaixo seguem os objetivos com dados no periodo."
                            )

                    df_plot = (
                        df_obj.groupby("date_pd", as_index=False)
                        .agg({"independent_rate": "mean"})
                        .sort_values("date_pd")
                    )
                    if df_plot.empty:
                        doc.add_paragraph("Sem dados numericos suficientes para gerar o grafico do objetivo.")
                        return

                    plt.figure(figsize=(7, 4))
                    plt.plot(df_plot["date_pd"], df_plot["independent_rate"], marker="o", label="Independencia")
                    plt.title(limpar_texto_pei(titulo)[:90])
                    plt.ylabel("Percentual")
                    plt.ylim(0, 100)
                    ax = plt.gca()
                    ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
                    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
                    plt.gcf().autofmt_xdate(rotation=45)
                    plt.legend()
                    plt.tight_layout()
                    buf_obj = io.BytesIO()
                    plt.savefig(buf_obj, format="png")
                    plt.close()
                    buf_obj.seek(0)
                    doc.add_picture(buf_obj, width=Inches(5.7))

                    df_barras = (
                        df_periodo.groupby("programa", as_index=False)["independent_rate"]
                        .mean()
                        .sort_values("independent_rate", ascending=True)
                        .tail(15)
                    )
                    if not df_barras.empty:
                        plt.figure(figsize=(7, 5))
                        plt.barh(df_barras["programa"], df_barras["independent_rate"], color="#2E86C1")
                        plt.title("Media de independencia por objetivo no periodo")
                        plt.xlabel("Independencia media (%)")
                        plt.xlim(0, 100)
                        plt.tight_layout()
                        buf_bar = io.BytesIO()
                        plt.savefig(buf_bar, format="png", dpi=160)
                        plt.close()
                        buf_bar.seek(0)
                        doc.add_picture(buf_bar, width=Inches(5.8))

                def adicionar_graficos_interferentes_periodo(doc, df_beh, inicio, fim_exclusivo):
                    doc.add_heading("Graficos de comportamentos interferentes", level=2)
                    df_beh_periodo = filtrar_periodo_pei(df_beh, inicio, fim_exclusivo)
                    if df_beh_periodo.empty:
                        doc.add_paragraph("Sem comportamentos interferentes registrados no periodo selecionado.")
                        return

                    if "count" in df_beh_periodo.columns:
                        df_totais = df_beh_periodo.groupby("comportamento")["count"].sum().reset_index()
                        if not df_totais.empty and df_totais["count"].sum() > 0:
                            plt.figure(figsize=(6.5, 4))
                            plt.bar(df_totais["comportamento"], df_totais["count"], color="#E74C3C")
                            plt.title("Frequencia total no periodo")
                            plt.xticks(rotation=45, ha="right")
                            plt.tight_layout()
                            buf_tot = io.BytesIO()
                            plt.savefig(buf_tot, format="png")
                            plt.close()
                            buf_tot.seek(0)
                            doc.add_picture(buf_tot, width=Inches(5.3))

                    if "rate" in df_beh_periodo.columns:
                        df_taxa = df_beh_periodo.groupby("comportamento")["rate"].mean().reset_index()
                        if not df_taxa.empty and df_taxa["rate"].sum() > 0:
                            plt.figure(figsize=(6.5, 4))
                            plt.bar(df_taxa["comportamento"], df_taxa["rate"], color="#F39C12")
                            plt.title("Taxa media no periodo")
                            plt.xticks(rotation=45, ha="right")
                            plt.tight_layout()
                            buf_taxa = io.BytesIO()
                            plt.savefig(buf_taxa, format="png")
                            plt.close()
                            buf_taxa.seek(0)
                            doc.add_picture(buf_taxa, width=Inches(5.3))

                    plt.figure(figsize=(7, 4))
                    for comportamento in df_beh_periodo["comportamento"].dropna().unique():
                        df_c = df_beh_periodo[df_beh_periodo["comportamento"] == comportamento].sort_values("date_pd")
                        y_col = "rate" if "rate" in df_c.columns else "count"
                        plt.plot(df_c["date_pd"], df_c[y_col], marker="o", label=str(comportamento))
                    plt.title("Evolucao dos interferentes no periodo")
                    ax = plt.gca()
                    ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
                    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
                    plt.gcf().autofmt_xdate(rotation=45)
                    plt.legend(bbox_to_anchor=(1.05, 1), loc="upper left")
                    plt.tight_layout()
                    buf_linha = io.BytesIO()
                    plt.savefig(buf_linha, format="png")
                    plt.close()
                    buf_linha.seek(0)
                    doc.add_picture(buf_linha, width=Inches(5.7))

                def adicionar_graficos_modelo(doc, df_hist, df_alvos, df_beh, objetivo_grafico, periodo_inicio, periodo_fim_exclusivo):
                    doc.add_page_break()
                    doc.add_heading("ANEXOS GRÁFICOS", level=1)
                    adicionar_grafico_objetivo_periodo(doc, df_hist, objetivo_grafico, periodo_inicio, periodo_fim_exclusivo)
                    adicionar_graficos_interferentes_periodo(doc, df_beh, periodo_inicio, periodo_fim_exclusivo)

                    if not df_hist.empty:
                        doc.add_heading("Resumo geral dos programas", level=2)
                        df_graf = df_hist.copy()
                        df_graf["date_pd"] = pd.to_datetime(df_graf["date"], errors="coerce")
                        df_graf = df_graf.dropna(subset=["date_pd"])
                        if not df_graf.empty:
                            df_linha = (
                                df_graf.groupby("date_pd", as_index=False)
                                .agg({"independent_rate": "mean", "prompt_rate": "mean"})
                                .sort_values("date_pd")
                            )
                            if not df_linha.empty:
                                plt.figure(figsize=(7, 4))
                                plt.plot(df_linha["date_pd"], df_linha["independent_rate"], marker="o", label="Independência média")
                                plt.plot(df_linha["date_pd"], df_linha["prompt_rate"], marker="o", label="Ajuda média")
                                plt.title("Evolução média dos programas")
                                plt.ylabel("Percentual")
                                plt.ylim(0, 100)
                                ax = plt.gca()
                                ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
                                ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
                                plt.gcf().autofmt_xdate(rotation=45)
                                plt.legend()
                                plt.tight_layout()
                                buf = io.BytesIO()
                                plt.savefig(buf, format="png")
                                plt.close()
                                buf.seek(0)
                                doc.add_picture(buf, width=Inches(5.7))

                    if not df_beh.empty:
                        df_beh_g = df_beh.copy()
                        df_beh_g["date_pd"] = pd.to_datetime(df_beh_g["date"], errors="coerce")
                        df_beh_g = df_beh_g.dropna(subset=["date_pd"])
                        if not df_beh_g.empty:
                            doc.add_heading("Comportamentos interferentes", level=2)
                            df_totais = df_beh_g.groupby("comportamento")["count"].sum().reset_index()
                            if not df_totais.empty:
                                plt.figure(figsize=(6, 4))
                                plt.bar(df_totais["comportamento"], df_totais["count"], color="#E74C3C")
                                plt.title("Frequência total de ocorrências")
                                plt.xticks(rotation=45, ha="right")
                                plt.tight_layout()
                                buf_tot = io.BytesIO()
                                plt.savefig(buf_tot, format="png")
                                plt.close()
                                buf_tot.seek(0)
                                doc.add_picture(buf_tot, width=Inches(5.2))

                            if "rate" in df_beh_g.columns:
                                df_taxa = df_beh_g.groupby("comportamento")["rate"].mean().reset_index()
                                if not df_taxa.empty:
                                    plt.figure(figsize=(6, 4))
                                    plt.bar(df_taxa["comportamento"], df_taxa["rate"], color="#F39C12")
                                    plt.title("Taxa média")
                                    plt.xticks(rotation=45, ha="right")
                                    plt.tight_layout()
                                    buf_taxa = io.BytesIO()
                                    plt.savefig(buf_taxa, format="png")
                                    plt.close()
                                    buf_taxa.seek(0)
                                    doc.add_picture(buf_taxa, width=Inches(5.2))

                            plt.figure(figsize=(7, 4))
                            for comportamento in df_beh_g["comportamento"].dropna().unique():
                                df_c = df_beh_g[df_beh_g["comportamento"] == comportamento].sort_values("date_pd")
                                plt.plot(df_c["date_pd"], df_c["count"], marker="o", label=comportamento)
                            plt.title("Evolução ao longo do tempo")
                            ax = plt.gca()
                            ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
                            ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
                            plt.gcf().autofmt_xdate(rotation=45)
                            plt.legend(bbox_to_anchor=(1.05, 1), loc="upper left")
                            plt.tight_layout()
                            buf_linha = io.BytesIO()
                            plt.savefig(buf_linha, format="png")
                            plt.close()
                            buf_linha.seek(0)
                            doc.add_picture(buf_linha, width=Inches(5.7))

                def gerar_doc_modelo_pei(
                    nome_paciente,
                    df_prog,
                    df_hist,
                    df_alvos,
                    df_lib,
                    df_beh,
                    objetivo_grafico,
                    periodo_inicio,
                    periodo_fim,
                    periodo_fim_exclusivo,
                    texto_analise_trimestral,
                    texto_resumo_comportamentos,
                ):
                    doc = Document(PEI_TEMPLATE_PATH)
                    aplicar_fonte_pei(doc)
                    df_prog_periodo = filtrar_periodo_pei(df_prog, periodo_inicio, periodo_fim_exclusivo)
                    areas = distribuir_objetivos_por_area(df_prog_periodo, df_beh)
                    df_alvos_periodo = filtrar_periodo_pei(df_alvos, periodo_inicio, periodo_fim_exclusivo)

                    if doc.tables:
                        dados = (
                            "DADOS DE IDENTIFICAÇÃO\n"
                            f"Nome do(a) Aprendiz: {nome_paciente}\n"
                            "Data de Nascimento: N/I | Idade: N/I\n"
                            f"Data de Elaboração do Plano: {datetime.now().strftime('%m/%Y')}\n"
                            "Equipe Responsável: Sistema Skinner\n"
                            "Equipe Atual: Aplicável em caso de mudança de equipe\n"
                            "Instrumento(s) de Avaliação Utilizado(s): bHave, ABLLS-R e registros clínicos\n"
                            f"Data da Avaliação: {START_DATE.strftime('%d/%m/%Y')}\n"
                            f"Data da Finalização: {datetime.now().strftime('%d/%m/%Y')}"
                        )
                        set_cell_text(doc.tables[0].cell(1, 0), dados)

                    for area_num in range(1, 6):
                        table_index = area_num
                        if table_index < len(doc.tables):
                            preencher_tabela_objetivos(
                                doc.tables[table_index],
                                area_num,
                                areas.get(area_num, []),
                            )

                    desempenho_tables = {1: 6, 2: 7, 3: 8}
                    for area_num, table_index in desempenho_tables.items():
                        if table_index < len(doc.tables):
                            preencher_tabela_desempenho(
                                doc.tables[table_index],
                                area_num,
                                areas.get(area_num, []),
                                df_hist,
                                df_alvos,
                                df_alvos_periodo,
                                df_lib,
                                periodo_inicio,
                                periodo_fim,
                            )

                    adicionar_resumo_clinico_periodo(
                        doc,
                        texto_resumo_comportamentos,
                        periodo_inicio,
                        periodo_fim,
                    )
                    adicionar_graficos_modelo(doc, df_hist, df_alvos, df_beh, objetivo_grafico, periodo_inicio, periodo_fim_exclusivo)
                    aplicar_fonte_pei(doc)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    return buffer

                def gerar_doc_completo(
                    nome_paciente,
                    df_prog,
                    df_hist,
                    df_alvos,
                    df_lib,
                    df_beh,
                    objetivo_grafico,
                    periodo_inicio,
                    periodo_fim,
                    periodo_fim_exclusivo,
                    texto_analise_trimestral,
                    texto_resumo_comportamentos,
                ):
                    if os.path.exists(PEI_TEMPLATE_PATH):
                        return gerar_doc_modelo_pei(
                            nome_paciente,
                            df_prog,
                            df_hist,
                            df_alvos,
                            df_lib,
                            df_beh,
                            objetivo_grafico,
                            periodo_inicio,
                            periodo_fim,
                            periodo_fim_exclusivo,
                            texto_analise_trimestral,
                            texto_resumo_comportamentos,
                        )

                    doc = Document()
                    aplicar_fonte_pei(doc)
                    
                    status_objetivos = {"ATG": 0, "AVA": 0, "EST": 0, "AGR": 0, "-": 0}
                    total_alvos_trabalhados = 0
                    total_alvos_atingidos = 0
                    
                    title = doc.add_heading('PLANO DE ENSINO INDIVIDUALIZADO (PEI)', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_heading('DADOS DE IDENTIFICAÇÃO', level=1)
                    t_id = doc.add_table(rows=4, cols=2)
                    aplicar_estilo_tabela_seguro(t_id)
                    t_id.rows[0].cells[0].text = f"Nome do(a) Aprendiz: {nome_paciente}"
                    t_id.rows[1].cells[0].text = f"Data de Início da Análise: {START_DATE.strftime('%d/%m/%Y')}"
                    t_id.rows[2].cells[0].text = f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y')}"
                    t_id.rows[3].cells[0].text = "Equipe Responsável: Sistema Skinner"

                    df_alvos_periodo = filtrar_periodo_pei(df_alvos, periodo_inicio, periodo_fim_exclusivo)
                    adicionar_resumo_clinico_periodo(
                        doc,
                        texto_resumo_comportamentos,
                        periodo_inicio,
                        periodo_fim,
                    )

                    doc.add_heading('OBJETIVOS DE INTERVENÇÃO', level=1)
                    doc.add_paragraph("(As metas devem ser Específicas, Mensuráveis, Alcançáveis, Relevantes e com Prazo definido.)\n")

                    df_prog_unicos = df_prog.drop_duplicates(subset=['programa'])
                    df_prog_periodo = filtrar_periodo_pei(df_prog, periodo_inicio, periodo_fim_exclusivo)
                    areas = distribuir_objetivos_por_area(df_prog_periodo, df_beh)

                    for area_num in range(1, 6):
                        doc.add_heading(f"Area {area_num}", level=2)
                        tabela_area = doc.add_table(rows=1, cols=2)
                        preencher_tabela_objetivos(
                            tabela_area,
                            area_num,
                            areas.get(area_num, []),
                        )

                    doc.add_heading('ALVOS TRABALHADOS E DESEMPENHO', level=1)
                    for area_num in range(1, 6):
                        doc.add_heading(f"Area {area_num}", level=2)
                        tabela_alvos = doc.add_table(rows=18, cols=4)
                        preencher_tabela_desempenho(
                            tabela_alvos,
                            area_num,
                            areas.get(area_num, []),
                            df_hist,
                            df_alvos,
                            df_alvos_periodo,
                            df_lib,
                            periodo_inicio,
                            periodo_fim,
                        )

                    for index, row in pd.DataFrame().iterrows():
                        prog_nome = row['programa']
                        doc.add_heading(f"{prog_nome}", level=2)
                        
                        texto_objetivo = row.get('objective', '')
                        doc.add_paragraph(verificar_alvos_clean(texto_objetivo))
                        
                        if not df_alvos.empty and 'programa' in df_alvos.columns:
                            alvos_prog = df_alvos[df_alvos['programa'] == prog_nome]
                            if not alvos_prog.empty:
                                nomes_alvos = ", ".join(alvos_prog['target_name'].unique())
                                p_alvo = doc.add_paragraph(f"Alvos: {nomes_alvos}")
                                p_alvo.runs[0].italic = True

                    doc.add_page_break()

                    doc.add_heading('DESEMPENHO POR ÁREA', level=1)
                    
                    trimestres = [
                        (START_DATE, START_DATE + timedelta(days=90), "Trimestre 1"),
                        (START_DATE + timedelta(days=90), START_DATE + timedelta(days=180), "Trimestre 2"),
                        (START_DATE + timedelta(days=180), START_DATE + timedelta(days=270), "Trimestre 3"),
                        (START_DATE + timedelta(days=270), START_DATE + timedelta(days=360), "Trimestre 4")
                    ]

                    for index, row in df_prog_unicos.iterrows():
                        prog_nome = row['programa']
                        limiar = row.get('mastery_threshold_percent', 90)
                        if pd.isna(limiar) or not limiar: limiar = 90
                        
                        doc.add_heading(f"{prog_nome}", level=3)
                        
                        hist_p = df_hist[df_hist['programa'] == prog_nome].copy()
                        linha_de_base = None
                        media_int_geral = None
                        status_lb = ""
                        
                        if not hist_p.empty:
                            hist_p['date_pd'] = pd.to_datetime(hist_p['date'])
                            if 'phase' not in hist_p.columns: hist_p['phase'] = ""
                            
                            mask_hist_lb = hist_p['phase'].astype(str).str.contains(r'linha de base|\blb\b|baseline|sondagem', case=False, na=False)
                            fases_lb = hist_p[mask_hist_lb]
                            fases_int = hist_p[~mask_hist_lb]
                            
                            if not fases_lb.empty:
                                linha_de_base = fases_lb['independent_rate'].mean()
                                if linha_de_base >= limiar: status_lb = "Objetivo atingido ja na Linha de Base"
                                elif not fases_int.empty: status_lb = "Em Intervencao (avancou apos Linha de Base)"
                                else: status_lb = "Avaliacao em andamento (em Linha de Base)"
                            else:
                                linha_de_base = hist_p.iloc[0]['independent_rate']
                                status_lb = "Em Intervencao (sem registro formal de LB no periodo)"
                                
                            if not fases_int.empty:
                                media_int_geral = fases_int['independent_rate'].mean()
                        
                        p_lb = doc.add_paragraph()
                        p_lb.add_run("Média Linha de Base: ").bold = True
                        p_lb.add_run(f"{linha_de_base:.1f}%" if pd.notnull(linha_de_base) else "N/A")
                        
                        p_int = doc.add_paragraph()
                        p_int.add_run("Média Intervenção: ").bold = True
                        p_int.add_run(f"{media_int_geral:.1f}%" if pd.notnull(media_int_geral) else "N/A")
                        
                        if status_lb:
                            p_st = doc.add_paragraph(status_lb)
                            p_st.runs[0].italic = True

                        table = doc.add_table(rows=1, cols=3)
                        aplicar_estilo_tabela_seguro(table)
                        hdr = table.rows[0].cells
                        hdr[0].text = 'Avaliação - Período'; hdr[1].text = 'Desempenho'; hdr[2].text = 'Código de evolução'
                        for cell in hdr: cell.paragraphs[0].runs[0].bold = True
                        
                        desempenho_anterior = linha_de_base 
                        evo_final_obj = "-" 
                        
                        for ini, fim, rotulo in trimestres:
                            row_cells = table.add_row().cells
                            if not hist_p.empty:
                                df_tri = hist_p[(hist_p['date_pd'] >= pd.Timestamp(ini)) & (hist_p['date_pd'] < pd.Timestamp(fim))]
                                if not df_tri.empty:
                                    media = df_tri['independent_rate'].mean()
                                    cod_evo = calcular_evolucao_trimestral(desempenho_anterior, media, limiar)
                                    row_cells[0].text = rotulo
                                    row_cells[1].text = f"{media:.1f}%"
                                    row_cells[2].text = cod_evo
                                    desempenho_anterior = media
                                    
                                    if "(" in cod_evo: evo_final_obj = cod_evo.split("(")[1].replace(")", "")
                                    else: evo_final_obj = cod_evo
                                    continue
                            
                            row_cells[0].text = rotulo; row_cells[1].text = "-"; row_cells[2].text = "Sem avaliações"
                        
                        if evo_final_obj in status_objetivos: status_objetivos[evo_final_obj] += 1
                        else: status_objetivos["-"] += 1

                        if not df_alvos.empty and 'programa' in df_alvos.columns:
                            alvos_prog = df_alvos[df_alvos['programa'] == prog_nome].copy()
                            p_alvos = doc.add_paragraph()
                            p_alvos.add_run("\nAlvos que atingiram a meta no periodo:").bold = True
                            
                            if not alvos_prog.empty:
                                if 'phase' not in alvos_prog.columns: alvos_prog['phase'] = ""
                                total_alvos_trabalhados += alvos_prog['target_name'].nunique()
                                
                                lb_mastered = alvos_prog[alvos_prog['phase'].astype(str).str.contains(r'linha de base|\blb\b|baseline|sondagem', case=False, na=False) & (alvos_prog['independent_rate'] >= 100)]['target_name'].unique()
                                mask = ~alvos_prog['target_name'].isin(lb_mastered)
                                
                                if not alvos_prog[mask].empty:
                                    max_rates = alvos_prog[mask].groupby('target_name')['independent_rate'].max()
                                    alvos_meta = [(alvo, taxa) for alvo, taxa in max_rates.items() if taxa >= limiar]
                                    total_alvos_atingidos += len(alvos_meta)
                                    
                                    if alvos_meta:
                                        for alvo, taxa in sorted(alvos_meta):
                                            doc.add_paragraph(f"  - {alvo}: {taxa:.1f}% de independencia")
                                    else: doc.add_paragraph("  - Nenhum novo alvo atingiu o criterio de dominio.")
                                else: doc.add_paragraph("  - Nenhum novo alvo atingiu o criterio de dominio.")
                            else: doc.add_paragraph("  - Sem registros granulares de alvos no periodo.")
                        doc.add_paragraph()

                    # --- RESUMO GRÁFICO ---
                    doc.add_page_break()
                    doc.add_heading('RESUMO CLÍNICO DE EVOLUÇÃO', level=1)
                    adicionar_grafico_objetivo_periodo(doc, df_hist, objetivo_grafico, periodo_inicio, periodo_fim_exclusivo)
                    adicionar_graficos_interferentes_periodo(doc, df_beh, periodo_inicio, periodo_fim_exclusivo)
                    
                    labels_obj, sizes_obj, colors_obj = [], [], []
                    color_map = {"ATG": "#2ECC71", "AVA": "#3498DB", "EST": "#F1C40F", "AGR": "#E74C3C", "-": "#95A5A6"}
                    nome_map = {"ATG": "Atingiu", "AVA": "Avançou", "EST": "Estabilizou", "AGR": "Agravou", "-": "S/ Dados"}

                    for k, v in status_objetivos.items():
                        if v > 0:
                            labels_obj.append(nome_map.get(k, k))
                            sizes_obj.append(v)
                            colors_obj.append(color_map.get(k, "#333333"))

                    if sizes_obj:
                        plt.figure(figsize=(5, 4))
                        plt.pie(sizes_obj, labels=labels_obj, colors=colors_obj, autopct='%1.1f%%', startangle=140)
                        plt.title('Status Final dos Programas')
                        plt.tight_layout()
                        buf_obj = io.BytesIO()
                        plt.savefig(buf_obj, format='png')
                        plt.close()
                        buf_obj.seek(0)
                        doc.add_picture(buf_obj, width=Inches(4.5))

                    if total_alvos_trabalhados > 0:
                        plt.figure(figsize=(5, 4))
                        bars = plt.bar(['Trabalhados', 'Atingiram Meta'], [total_alvos_trabalhados, total_alvos_atingidos], color=['#3498DB', '#2ECC71'])
                        plt.title('Aquisição de Novos Alvos')
                        plt.ylabel('Quantidade de Alvos')
                        for bar in bars:
                            yval = bar.get_height()
                            plt.text(bar.get_x() + bar.get_width()/2, yval + 0.5, int(yval), ha='center', va='bottom', fontweight='bold')
                        plt.tight_layout()
                        buf_alvos = io.BytesIO()
                        plt.savefig(buf_alvos, format='png')
                        plt.close()
                        buf_alvos.seek(0)
                        doc.add_picture(buf_alvos, width=Inches(4.5))

                    if not df_hist.empty:
                        df_graf = df_hist.copy()
                        df_graf["date_pd"] = pd.to_datetime(df_graf["date"], errors="coerce")
                        df_graf = df_graf.dropna(subset=["date_pd"])
                        if not df_graf.empty:
                            doc.add_heading("Graficos de desempenho dos programas", level=2)

                            df_linha = (
                                df_graf.groupby("date_pd", as_index=False)
                                .agg({"independent_rate": "mean", "prompt_rate": "mean"})
                                .sort_values("date_pd")
                            )
                            if not df_linha.empty:
                                plt.figure(figsize=(7, 4))
                                plt.plot(df_linha["date_pd"], df_linha["independent_rate"], marker="o", label="Independencia media")
                                plt.plot(df_linha["date_pd"], df_linha["prompt_rate"], marker="o", label="Ajuda media")
                                plt.title("Evolucao media dos programas")
                                plt.ylabel("Percentual")
                                plt.ylim(0, 100)
                                ax = plt.gca()
                                ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
                                ax.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m/%y"))
                                plt.gcf().autofmt_xdate(rotation=45)
                                plt.legend()
                                plt.tight_layout()
                                buf_prog_linha = io.BytesIO()
                                plt.savefig(buf_prog_linha, format="png")
                                plt.close()
                                buf_prog_linha.seek(0)
                                doc.add_picture(buf_prog_linha, width=Inches(5.5))

                            df_programas = (
                                df_graf.groupby("programa", as_index=False)["independent_rate"]
                                .mean()
                                .sort_values("independent_rate", ascending=True)
                                .tail(15)
                            )
                            if not df_programas.empty:
                                plt.figure(figsize=(7, 5))
                                plt.barh(df_programas["programa"], df_programas["independent_rate"], color="#2E86C1")
                                plt.title("Media de independencia por programa")
                                plt.xlabel("Percentual medio")
                                plt.xlim(0, 100)
                                plt.tight_layout()
                                buf_prog_bar = io.BytesIO()
                                plt.savefig(buf_prog_bar, format="png")
                                plt.close()
                                buf_prog_bar.seek(0)
                                doc.add_picture(buf_prog_bar, width=Inches(5.8))

                    if not df_beh.empty:
                        df_beh['date_pd'] = pd.to_datetime(df_beh['date'])
                        df_beh_pei = df_beh[df_beh['date_pd'] >= pd.Timestamp(START_DATE)]
                        
                        if not df_beh_pei.empty:
                            doc.add_page_break()
                            doc.add_heading('ANÁLISE DE COMPORTAMENTOS INTERFERENTES', level=1)

                            df_totais = df_beh_pei.groupby('comportamento')['count'].sum().reset_index()
                            if not df_totais.empty and df_totais['count'].sum() > 0:
                                plt.figure(figsize=(6, 4))
                                plt.bar(df_totais['comportamento'], df_totais['count'], color='#E74C3C')
                                plt.title('1. Frequência Total de Ocorrências')
                                plt.xticks(rotation=45, ha='right')
                                plt.tight_layout()
                                buf_tot = io.BytesIO()
                                plt.savefig(buf_tot, format='png')
                                plt.close()
                                buf_tot.seek(0)
                                doc.add_picture(buf_tot, width=Inches(5))

                            if 'rate' in df_beh_pei.columns:
                                df_taxa = df_beh_pei.groupby('comportamento')['rate'].mean().reset_index()
                                if not df_taxa.empty and df_taxa['rate'].sum() > 0:
                                    plt.figure(figsize=(6, 4))
                                    plt.bar(df_taxa['comportamento'], df_taxa['rate'], color='#F39C12')
                                    plt.title('2. Taxa Média (Ocorrências / Hora)')
                                    plt.xticks(rotation=45, ha='right')
                                    plt.tight_layout()
                                    buf_tx = io.BytesIO()
                                    plt.savefig(buf_tx, format='png')
                                    plt.close()
                                    buf_tx.seek(0)
                                    doc.add_picture(buf_tx, width=Inches(5))

                            plt.figure(figsize=(7, 4))
                            for comp in df_beh_pei['comportamento'].unique():
                                df_c = df_beh_pei[df_beh_pei['comportamento'] == comp].sort_values('date_pd')
                                plt.plot(df_c['date_pd'], df_c['count'], marker='o', label=comp)
                            
                            plt.title('3. Evolução ao Longo do Tempo')
                            ax = plt.gca()
                            ax.xaxis.set_major_locator(ticker.MaxNLocator(nbins=6))
                            ax.xaxis.set_major_formatter(mdates.DateFormatter('%d/%m/%y'))
                            plt.gcf().autofmt_xdate(rotation=45)
                            plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
                            plt.tight_layout()
                            buf_line = io.BytesIO()
                            plt.savefig(buf_line, format='png')
                            plt.close()
                            buf_line.seek(0)
                            doc.add_picture(buf_line, width=Inches(5.5))

                            col_tempo = next((col for col in ['duration', 'tempo', 'time'] if col in df_beh_pei.columns), None)
                            if col_tempo and df_beh_pei[col_tempo].sum() > 0:
                                df_tempo = df_beh_pei.groupby('comportamento')[col_tempo].sum().reset_index()
                                plt.figure(figsize=(6, 4))
                                plt.bar(df_tempo['comportamento'], df_tempo[col_tempo], color='#8E44AD')
                                plt.title('4. Duração Total (Tempo Acumulado)')
                                plt.xticks(rotation=45, ha='right')
                                plt.tight_layout()
                                buf_tmp = io.BytesIO()
                                plt.savefig(buf_tmp, format='png')
                                plt.close()
                                buf_tmp.seek(0)
                                doc.add_picture(buf_tmp, width=Inches(5))

                    doc.add_page_break()
                    doc.add_heading('Legenda - Código de evolução', level=2)
                    legendas = [
                        ("Atingiu (ATG):", "O objetivo atingiu o critério de domínio no período avaliado."),
                        ("Avançou (AVA):", "Houve progresso significativo (> 5%) em relação à avaliação ou período anterior."),
                        ("Estabilizou (EST):", "Desempenho manteve-se na mesma margem do último período."),
                        ("Agravou (AGR):", "Houve perda de habilidade ou queda de desempenho no período avaliado.")
                    ]
                    for sigla, desc in legendas:
                        p = doc.add_paragraph()
                        p.add_run(sigla).bold = True
                        p.add_run(f" {desc}")

                    aplicar_fonte_pei(doc)
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    return buffer

                # Visualização na Tela do Dashboard
                st.subheader("Areas e objetivos no periodo selecionado")
                df_pei = filtrar_periodo_pei(df_p_raw, ciclo_inicio_pei, ciclo_fim_exclusivo_pei)
                df_alvos_preview = carregar_alvos_programas_pei(programas_pei)
                df_alvos_periodo_preview = filtrar_periodo_pei(df_alvos_preview, ciclo_inicio_pei, ciclo_fim_exclusivo_pei)
                areas_preview = distribuir_objetivos_por_area(df_pei, filtrar_periodo_pei(df_b_raw, ciclo_inicio_pei, ciclo_fim_exclusivo_pei))
                resumo_objetivos_preview = resumo_objetivos_por_area(areas_preview)
                if resumo_objetivos_preview.empty:
                    st.info("Sem objetivos registrados no periodo selecionado.")
                else:
                    st.dataframe(resumo_objetivos_preview, hide_index=True, use_container_width=True)

                st.subheader("Alvos trabalhados e desempenho no periodo selecionado")
                resumo_alvos_preview = resumo_alvos_por_objetivo(
                    areas_preview,
                    df_p_raw,
                    df_alvos_preview,
                    df_alvos_periodo_preview,
                    df_lib,
                    ciclo_inicio_pei,
                    ciclo_fim_pei,
                )
                if resumo_alvos_preview.empty:
                    st.info("Sem alvos trabalhados no periodo selecionado.")
                else:
                    st.dataframe(
                        estilizar_status_preview_pei(resumo_alvos_preview),
                        hide_index=True,
                        use_container_width=True,
                    )

                st.subheader("Grafico do objetivo selecionado")
                df_obj_visual = df_pei.copy()
                if objetivo_grafico_pei != "Media geral dos objetivos" and not df_obj_visual.empty:
                    df_obj_visual = df_obj_visual[df_obj_visual["programa"] == objetivo_grafico_pei]
                if df_obj_visual.empty:
                    st.info("Sem registros para o objetivo selecionado no periodo.")
                else:
                    df_obj_plot = (
                        df_obj_visual.groupby("date_pd", as_index=False)
                        .agg({"independent_rate": "mean"})
                        .sort_values("date_pd")
                    )
                    df_obj_plot = df_obj_plot.rename(columns={
                        "date_pd": "Data",
                        "independent_rate": "Independencia",
                    })
                    fig_obj_pei = px.line(
                        df_obj_plot,
                        x="Data",
                        y="Independencia",
                        markers=True,
                        title=objetivo_grafico_pei,
                        labels={"value": "Percentual", "variable": "Medida"},
                    )
                    fig_obj_pei.update_yaxes(range=[0, 100])
                    fig_obj_pei.update_xaxes(type="date", tickformat="%d/%m/%Y", tickangle=-45)
                    st.plotly_chart(fig_obj_pei, use_container_width=True)

                if not df_b_raw.empty:
                    st.subheader("Evolucao de Comportamentos Interferentes")
                    df_b_pei = filtrar_periodo_pei(df_b_raw, ciclo_inicio_pei, ciclo_fim_exclusivo_pei)
                    if not df_b_pei.empty:
                        y_interf = 'rate' if 'rate' in df_b_pei.columns else 'count'
                        fig_pei = px.bar(df_b_pei, x='date', y=y_interf, color='comportamento', title="Interferentes no periodo selecionado")
                        fig_pei.update_xaxes(type='date', tickformat="%d/%m/%Y", tickangle=-45)
                        st.plotly_chart(fig_pei, use_container_width=True)

                pei_key = f"pei_docx::{paciente_sel}"
                pei_name_key = f"pei_docx_name::{paciente_sel}"

                if st.button("Gerar Documento PEI Oficial com Graficos"):
                    with st.spinner("Compilando graficos, grade de objetivos e resumo dos comportamentos. Aguarde..."):
                        df_alvos_completo = df_alvos_preview.copy()
                        
                        df_prog_dados = df_p_raw.copy()
                        
                        # Limpeza do objetivo
                        if 'objective' in df_prog_dados.columns:
                            df_prog_dados['objective'] = df_prog_dados['objective'].astype(str).replace({'None': None, 'nan': None, '': None})
                        else:
                            df_prog_dados['objective'] = None
                            
                        # Merge lógico: Verifica se a biblioteca existe ANTES de usar o else
                        if not df_lib.empty and 'name' in df_lib.columns:
                            df_lib_unique = df_lib.drop_duplicates(subset=['name'])
                            df_prog_dados = df_prog_dados.merge(df_lib_unique[['name', 'mastery_threshold_percent', 'objective_template']], left_on='programa', right_on='name', how='left')
                            # Preenche o objetivo com o template da biblioteca se o bHave estiver vazio
                            df_prog_dados['objective'] = df_prog_dados['objective'].fillna(df_prog_dados['objective_template']).fillna("Descrição não informada.")
                        else:
                            # Se não houver biblioteca, preenche com o padrão
                            df_prog_dados['objective'] = df_prog_dados['objective'].fillna("Descrição não informada.")
                        
                        df_prog_periodo_pei = filtrar_periodo_pei(df_prog_dados, ciclo_inicio_pei, ciclo_fim_exclusivo_pei)
                        df_alvos_periodo_pei = filtrar_periodo_pei(df_alvos_completo, ciclo_inicio_pei, ciclo_fim_exclusivo_pei)
                        df_beh_periodo_pei = filtrar_periodo_pei(df_b_raw, ciclo_inicio_pei, ciclo_fim_exclusivo_pei)
                        texto_analise_trimestral = ""
                        texto_resumo_comportamentos = gerar_resumo_comportamentos_problema_pei(
                            paciente_sel,
                            df_beh_periodo_pei,
                            ciclo_inicio_pei,
                            ciclo_fim_pei,
                        )

                        # Gera o buffer do documento
                        buffer_file = gerar_doc_completo(
                            paciente_sel,
                            df_prog_dados,
                            df_p_raw,
                            df_alvos_completo,
                            df_lib,
                            df_b_raw,
                            objetivo_grafico_pei,
                            ciclo_inicio_pei,
                            ciclo_fim_pei,
                            ciclo_fim_exclusivo_pei,
                            texto_analise_trimestral,
                            texto_resumo_comportamentos,
                        )
                        st.session_state[pei_key] = buffer_file.getvalue()
                        st.session_state[pei_name_key] = (
                            f"PEI_{paciente_sel.replace(' ', '_')}_"
                            f"{ciclo_inicio_pei.strftime('%Y%m%d')}_{ciclo_fim_pei.strftime('%Y%m%d')}.docx"
                        )
                        st.success("PEI gerado. Use o botao abaixo para baixar o arquivo.")

                if st.session_state.get(pei_key):
                    st.download_button(
                        label="Baixar PEI em DOCX",
                        data=st.session_state[pei_key],
                        file_name=st.session_state.get(pei_name_key, f"PEI_{paciente_sel.replace(' ', '_')}.docx"),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

            # ==========================================
            # --- ABA 5: AVALIAÇÕES (ABLLS-R) ---
            # ==========================================
            with tab5:
                render_ablls_module(paciente_sel, df_p_raw)

            # ==========================================
            # --- ABA 6: ASSISTENTE ---
            # ==========================================
            with tab6:
                st.markdown(f"### Assistente clinico: **{paciente_sel}**")

                min_d_a, max_d_a = df_p_raw['date'].min(), df_p_raw['date'].max()
                c_dt_a, _ = st.columns([1, 2])
                periodo_a = c_dt_a.date_input(
                    "Periodo de analise",
                    value=(min_d_a, max_d_a),
                    format="DD/MM/YYYY",
                    key="assistente_periodo",
                )
                start_a, end_a = (periodo_a[0], periodo_a[1]) if len(periodo_a) == 2 else (min_d_a, max_d_a)

                resposta_agente_key = f"resposta_agente_clinico_{paciente_sel}"
                pergunta_agente = st.text_area(
                    "Pergunta clinica",
                    value=(
                        "Quais hipoteses clinicas e proximos passos observaveis fazem sentido "
                        "para este paciente neste periodo, considerando habilidades, alvos e comportamentos interferentes?"
                    ),
                    height=140,
                    key="pergunta_agente_clinico",
                )

                if st.button("Consultar assistente", key="botao_agente_clinico"):
                    if not pergunta_agente.strip():
                        st.warning("Informe uma pergunta clinica.")
                    else:
                        with st.spinner("Consultando a base local e preparando a sintese..."):
                            try:
                                st.session_state[resposta_agente_key] = ask_clinical_agent(
                                    paciente_sel,
                                    pergunta_agente,
                                    start_a,
                                    end_a,
                                )
                            except Exception as exc:
                                st.error(f"Erro ao consultar o assistente: {exc}")

                resposta_agente = st.session_state.get(resposta_agente_key)
                if resposta_agente:
                    st.markdown("#### Sintese")
                    modo_agente = resposta_agente.get("modo", "")
                    texto_resposta_agente = resposta_agente.get("resposta", "")
                    if modo_agente == "limite_ia":
                        st.warning(texto_resposta_agente)
                    elif modo_agente in ["busca_local", "busca_local_fallback"]:
                        st.info(texto_resposta_agente)
                    else:
                        st.write(texto_resposta_agente)
                    texto_referencias = texto_resposta_agente.casefold()
                    if (
                        resposta_agente.get("fontes")
                        and "referencias utilizadas" not in texto_referencias
                        and "referências utilizadas" not in texto_referencias
                    ):
                        render_agent_references(resposta_agente.get("fontes", []))
