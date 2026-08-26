from pathlib import Path

from docx import Document


path = Path(r"C:/Users/coord/Desktop/PEI_2026_TEMPLATE_V3.docx")
doc = Document(str(path))

print("paragraphs", len(doc.paragraphs))
print("tables", len(doc.tables))
print("sections", len(doc.sections))
print("header paragraphs", sum(len(section.header.paragraphs) for section in doc.sections))
print("footer paragraphs", sum(len(section.footer.paragraphs) for section in doc.sections))

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip().replace("\n", " ")
    if text:
        print(f"P{i}: {text[:180]}")

for section_index, section in enumerate(doc.sections):
    print(f"HEADER {section_index}:", [p.text.strip() for p in section.header.paragraphs if p.text.strip()])
    print(f"FOOTER {section_index}:", [p.text.strip() for p in section.footer.paragraphs if p.text.strip()])

for table_index, table in enumerate(doc.tables):
    print(f"TABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)}")
    max_rows = min(len(table.rows), 20)
    for row_index, row in enumerate(table.rows[:max_rows]):
        cells = [cell.text.strip().replace("\n", " | ")[:90] for cell in row.cells]
        print("  R", row_index, cells)

print("TABLE0 FULL")
print(doc.tables[0].cell(1, 0).text)
